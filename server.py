
from flask import Flask, render_template, request, jsonify, redirect, url_for, session
from flask_sqlalchemy import SQLAlchemy
import os
import re
import logging
from flask_cors import CORS
# from universal_job_scraper import main as run_scraper
from werkzeug.security import generate_password_hash, check_password_hash
import pdfplumber
import secrets
from datetime import timedelta
from datetime import datetime
import threading
from flask import render_template, session, redirect
from werkzeug.utils import secure_filename

from models import db, User, Candidate, Job, Company  # ✅ Import from models.py
from universal_job_scraper import main, save_jobs_to_db  # ✅ Now safe to import
from flask import Flask, request, jsonify, render_template, redirect, session

from flask_migrate import Migrate
from flask import Flask, session
from flask_session import Session
from flask_mail import Mail, Message
from indeed import run_indeed_for_candidate
import pandas as pd
import subprocess
import sys
import json
from openai import OpenAI
import pytesseract
from PIL import Image
from docx import Document  # For .docx
import mammoth  
from google_auth_oauthlib.flow import Flow
# import boto3
import uuid
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.date import DateTrigger
import boto3
from botocore.client import Config
from flask import jsonify
import threading
from werkzeug.middleware.proxy_fix import ProxyFix
from dotenv import load_dotenv
load_dotenv()

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CONTENT_DIR = os.path.join(BASE_DIR, "content")     # put your .txt files here

PRIVACY_FILE = os.path.join(CONTENT_DIR, "privacy_policy_text.txt")
TERMS_FILE   = os.path.join(CONTENT_DIR, "terms_and_condition.txt")


from dotenv import load_dotenv
load_dotenv()

app = Flask(__name__, template_folder='templates')

# ✅ Trust DigitalOcean’s proxy (so HTTPS is recognized)
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

# ✅ Optional: make url_for generate https:// links
app.config.update(PREFERRED_URL_SCHEME='https')

# --- Database config (hardcoded as requested) ---
app.config['SQLALCHEMY_DATABASE_URI'] = (
    ""
    ""
)
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# --- Session config ---
app.config['SESSION_TYPE'] = 'sqlalchemy'
app.config['SESSION_SQLALCHEMY'] = db

# --- Initialize DB ---
db.init_app(app)

# Ensure all tables (including 'sessions') exist before starting
with app.app_context():
    db.create_all()

# --- Initialize session ---
Session(app)
migrate = Migrate(app, db)

# Create DB tables
with app.app_context():
    db.create_all()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
app.logger.setLevel(logging.INFO)

# Directory to store resumes
RESUME_DIR = "resumes"
if not os.path.exists(RESUME_DIR):
    os.makedirs(RESUME_DIR)

# ------------------- S3 CONFIG -------------------
# S3_BUCKET = "resumes-folder"
# S3_REGION = "ap-south-1"
# s3_client = boto3.client("s3", region_name=S3_REGION)


# Configure your DigitalOcean Spaces info here or from environment variables
DO_SPACE_NAME = "resumes-folder"        # Your Spaces bucket name
DO_REGION = "blr1"                      # Your Spaces region, e.g., blr1, nyc3, sfo2 etc.
DO_ENDPOINT = f"https://{DO_REGION}.digitaloceanspaces.com"

DO_ACCESS_KEY = ""
DO_SECRET_KEY = ""

# Create boto3 client for DigitalOcean Spaces
boto3_session = boto3.session.Session()
spaces_client = boto3_session.client(
    's3',
    region_name=DO_REGION,
    endpoint_url=DO_ENDPOINT,
    aws_access_key_id=DO_ACCESS_KEY,
    aws_secret_access_key=DO_SECRET_KEY,
    config=Config(signature_version='s3v4')
)


def save_resume_to_spaces(file_obj, original_filename):
    ext = os.path.splitext(original_filename)[1]
    unique_suffix = uuid.uuid4().hex[:6]
    filename = f"{os.path.splitext(original_filename)[0]}_{unique_suffix}{ext}"
    space_key = f"resumes/{filename}"
    logging.info(f"[INFO] Uploading resume to DigitalOcean Spaces: {space_key}")
    app.logger.info(f"[INFO] Uploading resume to DigitalOcean Spaces: {space_key}")
    spaces_client.upload_fileobj(file_obj, DO_SPACE_NAME, space_key, ExtraArgs={'ACL': 'private'}) # Store privately
    return f"s3://{DO_SPACE_NAME}/{space_key}"

def get_resume_from_spaces(s3_path, local_temp_path):
    space_key = s3_path.replace(f"s3://{DO_SPACE_NAME}/", "")
    logging.info(f"[INFO] Downloading resume from DigitalOcean Spaces: {space_key}")
    app.logger.info(f"[INFO] Downloading resume from DigitalOcean Spaces: {space_key}")
    spaces_client.download_file(DO_SPACE_NAME, space_key, local_temp_path)
    return local_temp_path

# Route to view resume securely
@app.route('/-bkp-view-resume/<int:candidate_id>')
def view_resume_bkp(candidate_id):
    logging.info("Enter view_resume function.")
    app.logger.info("Enter view_resume function.")
    """
    Generates a temporary, secure URL to view a private resume from Spaces.
    """
    if 'user_id' not in session:
        return "Unauthorized", 401

    candidate = Candidate.query.get(candidate_id)
    logging.info(f"Inside view_resume function. candidate : {candidate}")
    app.logger.info(f"Inside view_resume function. candidate : {candidate}")
    if not candidate or not candidate.resume_path:
        return "Resume not found", 404

    # Extract the key from the s3:// path
    space_key = candidate.resume_path.replace(f"s3://{DO_SPACE_NAME}/", "")
    logging.info(f"Inside view_resume function. space_key : {space_key}")
    app.logger.info(f"Inside view_resume function. space_key : {space_key}")

    # Generate a presigned URL that is valid for 1 hour (3600 seconds)
    presigned_url = spaces_client.generate_presigned_url('get_object',
                                                         Params={'Bucket': DO_SPACE_NAME, 'Key': space_key},
                                                         ExpiresIn=3600)
    logging.info(f"Exit view_resume function. presigned_url : {presigned_url}")
    app.logger.info(f"Exit view_resume function. presigned_url : {presigned_url}")
    return redirect(presigned_url)



# ------------------- OPENAI CONFIG -------------------
OPENAI_API_KEY = ""
client = OpenAI(api_key=OPENAI_API_KEY)

app.secret_key = secrets.token_hex(16)
CORS(app, supports_credentials=True)

# Session config
app.config['SESSION_TYPE'] = 'sqlalchemy'
app.config['SESSION_SQLALCHEMY'] = db
app.config['SESSION_COOKIE_SAMESITE'] = "None"
app.config['SESSION_COOKIE_SECURE'] = True
app.config['SESSION_COOKIE_HTTPONLY'] = True

# Mail config
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'dr.nayakinaidu@gmail.com'
app.config['MAIL_PASSWORD'] = 'hqlk kfwy iomv djcc'
app.config['MAIL_DEFAULT_SENDER'] = 'dr.nayakinaidu@gmail.com'

mail = Mail(app)
migrate = Migrate(app, db)

with app.app_context():
    db.create_all()

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
app.logger.setLevel(logging.INFO)

# ------------------- ROUTES -------------------

@app.route('/')
def landing():
    return render_template('landing.html', api_token=os.getenv('API_TOKEN'),backend_url=os.getenv('BACKEND_URL'))
    
# ---------- OCR for scanned PDFs ----------
def extract_text_with_ocr(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            img = page.to_image(resolution=300).original
            text += pytesseract.image_to_string(Image.fromarray(img)) + "\n"
    return text

# ---------- Clean Word text ----------
def clean_word_text(text):
    text = re.sub(r"^[\s]*[\u2022\-\*\d]+\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    text = re.sub(r"\n?\s*\b\d+\b\s*\n?", "\n", text)
    return text.strip()

def extract_text_from_docx(docx_path):
    app.logger.info(f"[INFO] Entering extract_text_from_docx function")
    doc = Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Detect table rows separately
    education_table_rows = []
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                row_text = " | ".join(cells)
                tables_text.append(row_text)
                # Detect possible education rows
                if re.search(r"(B\.?Tech|B\.?E\.?|M\.?Tech|Diploma|High School|XII|College|University)",
                             row_text, re.IGNORECASE):
                    education_table_rows.append(row_text)

    full_text = "\n".join(paragraphs + tables_text)
    return clean_word_text(full_text), education_table_rows

def extract_text_from_doc(doc_path):
    with open(doc_path, "rb") as doc_file:
        result = mammoth.extract_raw_text(doc_file)
    return clean_word_text(result.value)

# ---------- Fallback section extraction ----------
# Add this mapping near the top of your file (or inside the module)
SECTION_VARIANTS = {
    "skills": [
        "skills", "skillset", "technical skills", "skills & tools", "software skills", "software tools"
    ],
    "experience": [
        "work experience", "experience", "professional experience", "employment history", "work history"
    ],
    "education": [
        "education", "educational qualifications", "educational qualification",
        "educational", "academics", "educational qualifications:", "educational qualification:"
    ],
    "certifications": [
        "certifications", "certificates", "licenses"
    ],
    "projects": [
        "projects", "personal projects", "academic projects", "project"
    ]
}

def _find_heading_end(text: str, variants: list) -> int:
    """
    Return index (end of match) of the first heading variant found (anchored at line start),
    or -1 if none.
    """
    for v in variants:
        # Anchor to start of line, allow optional whitespace and optional colon
        pattern = rf'(?im)^[\s]*{re.escape(v)}\b[\s:]*'
        m = re.search(pattern, text)
        if m:
            return m.end()
    return -1

def _find_next_heading_pos(text: str, start: int) -> int:
    """
    Find the earliest heading occurrence after `start`. Return its index or len(text) if none.
    """
    end = len(text)
    for variants in SECTION_VARIANTS.values():
        for v in variants:
            pattern = rf'(?im)^[\s]*{re.escape(v)}\b[\s:]*'
            m = re.search(pattern, text[start:], flags=re.IGNORECASE | re.MULTILINE)
            if m:
                pos = start + m.start()
                if pos < end:
                    end = pos
    return end

def extract_sections(text: str) -> dict:
    app.logger.info(f"[INFO] Entering extract_sections function")
    t = (text or "").strip()
    sections = {"key_skills": "", "employment": "", "education": "", "certifications": "", "projects": ""}

    # Helper to cut block for a given section using SECTION_VARIANTS anchors
    def cut_section(section_name: str) -> str:
        start = _find_heading_end(t, SECTION_VARIANTS.get(section_name, []))
        if start == -1:
            return ""
        end = _find_next_heading_pos(t, start)
        return t[start:end].strip()

    # Extract raw blocks using anchored headings
    raw_skills = cut_section("skills")
    raw_employment = cut_section("experience")
    raw_education = cut_section("education")
    raw_certifications = cut_section("certifications")
    raw_projects = cut_section("projects")

    # If education block is empty or looks malformed, attempt to extract table-like rows / degree lines
    if (not raw_education) or (len(raw_education) < 6) or re.match(r"^[a-z]{1,10}\s+qualification[:\s]*$", raw_education.strip().lower()):
        # find lines that look like education rows (table rows or degree lines)
        lines = [ln.strip() for ln in t.splitlines() if ln.strip()]
        degree_pattern = re.compile(
            r'\b(B\.?E\.?|BTECH|B\.?Tech|Bachelor|Bachelors|B\.?Sc|M\.?Tech|MTech|M\.?E\.?|MBA|XII|Xth|High School|Diploma)\b',
            flags=re.IGNORECASE
        )
        candidate_lines = []
        for ln in lines:
            # Accept table rows using '|' (from docx table export) or lines containing degree keywords or a 4-digit year
            if "|" in ln and re.search(r'\b(19|20)\d{2}\b', ln):
                candidate_lines.append(ln)
            elif degree_pattern.search(ln):
                candidate_lines.append(ln)
            elif re.search(r'\b(19|20)\d{2}\b', ln) and len(ln.split()) < 12:
                candidate_lines.append(ln)
        if candidate_lines:
            # join a few likely lines
            raw_education = "\n".join(candidate_lines[:10])

    # ----------------- Cleaning functions -----------------
    def clean_skills(raw: str) -> str:
        app.logger.info(f"[INFO] Entering clean_skills function")
        if not raw:
            return ""
        parts = re.split(r"[,;\n]", raw)
        stopwords = {"and","or","the","in","of","with","through","during","by","for","to","a","an","as"}
        bad_verbs = ["developed","designed","collaborated","created","delivered","implemented","achieved","contributed","adapted","enhanced","optimized","built","managed","worked","using","leveraging"]
        skills = []
        for p in parts:
            s = re.sub(r"\s+", " ", p).strip(" .,-")
            if not s or len(s) > 30:
                continue
            words = s.split()
            if len(words) > 4:
                continue
            if any(bv in s.lower() for bv in bad_verbs):
                continue
            if all(w.lower() in stopwords for w in words):
                continue
            skills.append(s)
        # dedupe preserving order
        seen = set()
        out = []
        for sk in skills:
            lk = sk.lower()
            if lk not in seen:
                seen.add(lk)
                out.append(sk)
        return ", ".join(out[:50])

    def clean_employment(raw: str) -> str:
        app.logger.info(f"[INFO] Entering clean_employment function")
        if not raw:
            return ""
        lines = [l.strip() for l in raw.splitlines() if l.strip()]
        date_pattern = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)?\s?\d{4}"
        cleaned = []
        for line in lines:
            line_clean = re.sub(r"^[\-\•\d\.\)\s]+", "", line)
            start, end = "", ""
            if "-" in line_clean and re.search(date_pattern, line_clean, flags=re.IGNORECASE):
                parts = re.split(r"\s*-\s*", line_clean)
                if len(parts) >= 2:
                    start = parts[0].strip()
                    end = parts[1].strip()
            comp_desig = re.sub(date_pattern, "", line_clean, flags=re.IGNORECASE).strip(" .,-")
            if comp_desig or start or end:
                pieces = []
                if comp_desig:
                    pieces.append(comp_desig)
                if start or end:
                    pieces.append(f"{start} - {end}".strip(" -"))
                cleaned.append(". ".join(pieces) if len(pieces) > 1 else pieces[0])
        return "\n".join(cleaned)

    def parse_education_block(raw: str) -> str:
        app.logger.info(f"[INFO] Entering parse_education_block function")
        if not raw:
            return ""
        lines = [l.strip() for l in raw.splitlines() if l.strip()]
        cleaned = []
        for line in lines:
            # If it's a table row with pipes, try to map cells to columns
            if "|" in line:
                parts = [p.strip() for p in line.split("|") if p.strip()]
                # Heuristic: if first cell is serial/no, then degree is second, inst third, board fourth, year last
                year = ""
                for part in reversed(parts):
                    m = re.search(r'\b(19|20)\d{2}\b', part)
                    if m:
                        year = m.group(0)
                        break
                degree = ""
                institution = ""
                # find institution by presence of 'college','institute','university','school'
                for p in parts:
                    if re.search(r'\b(college|institute|university|school|academy)\b', p, re.IGNORECASE):
                        institution = p
                        break
                # degree candidates
                for p in parts:
                    if re.search(r'\b(B\.?E\.?|BTECH|B\.?Tech|Bachelor|Bachelors|M\.?Tech|MTech|MBA|Diploma|XII|Xth)\b', p, re.IGNORECASE):
                        degree = p
                        break
                # fallback assignment if empty
                if not degree and len(parts) >= 2:
                    degree = parts[1]
                if not institution and len(parts) >= 3:
                    institution = parts[2]
                # Build string "Institution - Degree. Start - End" (we'll set end_date = year)
                entry = institution if institution else degree
                if degree and institution:
                    entry = f"{institution} - {degree}"
                if year:
                    entry += f". {year} - {year}"
                cleaned.append(entry.strip(" .,-"))
            else:
                # Non-table line: try extracting degree/institution/year with regex
                year_match = re.search(r'\b(19|20)\d{2}\b', line)
                year = year_match.group(0) if year_match else ""
                deg_match = re.search(r"((?:B\.?E\.?|BTECH|B\.?Tech|Bachelor(?:'s|\s)?|M\.?Tech|MTech|MBA|Diploma|XII|Xth|High School)[^,;\n]*)", line, re.IGNORECASE)
                degree = deg_match.group(1).strip() if deg_match else ""
                inst_match = re.search(r'(.+?)(?:,|\-| at | in | for | \| )', line)
                institution = ""
                if 'college' in line.lower() or 'institute' in line.lower() or 'university' in line.lower() or 'school' in line.lower():
                    # fallback: take the longer chunk not matching degree/year
                    parts = [p.strip() for p in re.split(r'[,;\-|]', line) if p.strip()]
                    for p in parts:
                        if degree and degree.lower() in p.lower():
                            continue
                        if re.search(r'\b(19|20)\d{2}\b', p):
                            continue
                        if len(p) > 3 and not re.search(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\b', p, re.IGNORECASE):
                            institution = p
                            break
                if not institution:
                    # fallback to using the remainder text after removing degree and year
                    tmp = line
                    if degree:
                        tmp = tmp.replace(degree, "")
                    if year:
                        tmp = tmp.replace(year, "")
                    tmp = re.sub(r'[\|\-\:,]+', ' ', tmp).strip()
                    if len(tmp) > 3:
                        institution = tmp
                entry = (institution or degree or line).strip()
                if year:
                    entry += f". {year} - {year}"
                cleaned.append(entry.strip(" .,-"))
        # dedupe
        seen = set()
        out = []
        for c in cleaned:
            lk = c.lower()
            if lk not in seen:
                seen.add(lk)
                out.append(c)
        return "\n".join(out)

    # build final cleaned sections
    sections["key_skills"] = clean_skills(raw_skills)
    sections["employment"] = clean_employment(raw_employment)
    sections["education"] = parse_education_block(raw_education)
    sections["certifications"] = (raw_certifications.replace("\n", ", ") or "").strip()
    sections["projects"] = "\n".join([ln.strip() for ln in raw_projects.splitlines() if ln.strip()]) if raw_projects else ""

    # last resort: if all empty, set skills to first chunk so UI doesn't break
    if not any(sections.values()):
        sections["key_skills"] = t[:2000]

    return sections

# ---------- Main parser ----------
def parse_resume(resume_path, use_llm=True):
    app.logger.info(f"[INFO] Entering parse_resume function")
    """
    Parse resume text from local path (already downloaded from S3).
    Runs regex first; calls LLM only if important fields are missing.
    """
    ext = os.path.splitext(resume_path)[1].lower()
    text = ""
    edu_rows = []

    try:
        if ext == ".pdf":
            with pdfplumber.open(resume_path) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            if not text.strip():
                print("[parse_resume] Using OCR for scanned PDF")
                app.logger.info(f"[INFO] parse_resume Using OCR for scanned PDF")
                text = extract_text_with_ocr(resume_path)
        elif ext == ".docx":
            text, edu_rows = extract_text_from_docx(resume_path)
        elif ext == ".doc":
            text = extract_text_from_doc(resume_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    except Exception as e:
        print(f"[parse_resume] Error reading file: {e}")
        app.logger.error(f"[parse_resume] Error reading file: {e}")
        return {k: "" for k in ["key_skills","employment","education","certifications","projects"]}

    regex_result = extract_sections(text)

    # If LLM not requested, return regex immediately
    if not use_llm:
        return regex_result

    # Decide if LLM is needed (fields missing or too short)
    needs_llm = (
        not regex_result.get("key_skills") or len(regex_result["key_skills"]) < 5 or
        not regex_result.get("employment") or len(regex_result["employment"]) < 20 or
        not regex_result.get("education") or len(regex_result["education"]) < 5
    )

    if not needs_llm:
        return regex_result

    llm_result = extract_sections_llm(text)
    return merge_parsed(regex_result, llm_result)

def extract_sections_llm(raw_text: str) -> dict:
    app.logger.info(f"[INFO] Enter extract_sections_llm function")
    """
    Calls OpenAI to extract key_skills, education, employment, certifications, and projects.
    Returns {} if API key missing or call fails.
    """
    api_key = os.getenv("OPENAI_API_KEY") or OPENAI_API_KEY
    if not api_key:
        return {}

    schema = {
        "type": "object",
        "properties": {
            "key_skills": {"type": "array", "items": {"type": "string"}},
            "education": {"type": "array", "items": {"type": "string"}},
            "employment": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "company": {"type": "string"},
                        "location": {"type": "string"},
                        "dates": {"type": "string"},
                        "title": {"type": "string"},
                        "highlights": {"type": "array", "items": {"type": "string"}}
                    }
                }
            },
            "certifications": {"type": "array", "items": {"type": "string"}},
            "projects": {"type": "array", "items": {"type": "string"}}
        }
    }

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "You are a precise information extractor. Output ONLY JSON matching the schema."},
                {"role": "user", "content": f"Extract key_skills, education, employment, certifications, and projects from this resume:\n<<<\n{raw_text}\n>>>\nSchema:\n{json.dumps(schema)}"}
            ],
            max_tokens=1500
        )
        return json.loads(resp.choices[0].message.content)
    except Exception as e:
        print(f"[extract_sections_llm] OpenAI extraction failed: {e}")
        app.logger.error(f"[extract_sections_llm] OpenAI extraction failed: {e}")
        return {}

def _norm_list_str(items):
    app.logger.info("Enter _norm_list_str function")
    dedup, seen = [], set()
    for s in items or []:
        s = re.sub(r"\s+", " ", s.strip())
        if s and s.lower() not in seen:
            seen.add(s.lower())
            dedup.append(s)
    return dedup

def merge_parsed(regex_out: dict, llm_out: dict) -> dict:
    app.logger.info("Enter merge_parsed function")
    """
    Merge regex and LLM outputs.
    Keep regex where it has data; fill blanks from LLM.
    """
    if not llm_out:
        return regex_out

    out = dict(regex_out)

    for k in ["key_skills", "education", "certifications", "projects"]:
        primary = _norm_list_str(regex_out.get(k, []))
        fallback = _norm_list_str(llm_out.get(k, []))
        out[k] = primary if primary else fallback

    regex_emp = regex_out.get("employment", []) or []
    llm_emp = llm_out.get("employment", []) or []
    out["employment"] = regex_emp if regex_emp else llm_emp

    return out

# Example usage
if __name__ == "__main__":
    resume_file = "Bhavesh Walankar Resume.pdf"
    parsed_data = parse_resume(resume_file)
    print(json.dumps(parsed_data, indent=2))

@app.route('/privacy')
def privacy():
    return render_template('privacy.html', api_token=os.getenv('API_TOKEN'), backend_url=os.getenv('BACKEND_URL'))  # uses templates/privacy.html

@app.route('/terms')
def terms():
    return render_template('terms.html',api_token=os.getenv('API_TOKEN'), backend_url=os.getenv('BACKEND_URL'))    # uses templates/terms.html



@app.route('/login')
def login():
    role = request.args.get('role')
    session['login_role'] = role
    return render_template('login_recruiter.html', api_token=os.getenv('API_TOKEN'),backend_url=os.getenv('BACKEND_URL')) if role == 'recruiter' else render_template('login.html', role=role, api_token=os.getenv('API_TOKEN'),backend_url=os.getenv('BACKEND_URL'))

@app.route('/login', methods=['POST'])
def login_submit():
    app.logger.info("Enter login_submit function")
    email = request.form.get('email')
    password = request.form.get('password')
    role = session.get('login_role')

    user = User.query.filter_by(email=email).first()
    if not user or not check_password_hash(user.password, password):
        return jsonify({'error': 'Invalid credentials'}), 401
    # Only allow login for approved users

    if not user.is_approved:
        return jsonify({'error': 'Your account is not approved yet.'}), 403
    
    # Ensure role matches
    # Allow jobseekers to log in from employee portal
    if role == 'employee' and user.role == 'jobseeker':
        pass  # allow

    elif user.role != role and not (user.role == 'admin' and role == 'recruiter'):
        return jsonify({'error': f'You are a {user.role}, not a {role}'}), 403
    
    session['user_id'] = user.id
    session['email'] = user.email
    session['role'] = user.role
        # 🔁 This is the part you replace
    # ✅ Fix this block:
    redirect_url = url_for('login_success')

    return jsonify({'message': 'Login successful','role':user.role, 'redirect': redirect_url})

@app.route('/signup', methods=['POST'])
def signup():
    try:
        data = request.form
        if User.query.filter_by(email=data['email']).first():
            return jsonify({'error': 'User already exists'}), 400

        company = Company(name=data['company_name'], location=data['company_location'])
        db.session.add(company)
        db.session.flush()

        new_user = User(
            first_name=data['first_name'],
            last_name=data['last_name'],
            email=data['email'],
            phone=data['phone'],
            password=generate_password_hash(data['password']),
            role='admin',  # ✅ Set admin role explicitly
            company_id=company.id,
            is_approved=False  # ✅ Recruiters require approval
        )
        db.session.add(new_user)
        db.session.flush()  # Ensure new_user.id is available

        # ✅ Send approval email to admin
        approval_token = os.getenv('API_TOKEN')
        approval_link = url_for('approve_recruiter', user_id=new_user.id, token=approval_token , _external=True)
        msg = Message(
            subject="New Recruiter Signup Approval Needed",
            recipients=['dr.nayakinaidu@gmail.com'],  # 🔁 Replace this
            body=f"A new recruiter ({new_user.email}) signed up.\n\nApprove them here: {approval_link}"
        )
        mail.send(msg)

        db.session.commit()

        # Don't set session — user should not be logged in yet
        return jsonify({'message': 'Signup successful. Your request is pending approval. You will be notified once approved.'})


    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

from werkzeug.security import generate_password_hash
import traceback

from google_auth_oauthlib.flow import Flow
import traceback

# --- Google OAuth Config ---
SCOPES = [
    'https://www.googleapis.com/auth/gmail.readonly',
    'https://www.googleapis.com/auth/gmail.send'
]
REDIRECT_URI = 'https://minute-gallery--automation-application-apify-backend.apify.actor/oauth2callback'

import os
from urllib.parse import urlencode, urlparse, urlunparse, parse_qsl

BASE_REDIRECT_URI = "https://minute-gallery--automation-application-apify-backend.apify.actor/oauth2callback"
API_TOKEN = os.getenv("API_TOKEN", "")

# Parse existing URL
url_parts = list(urlparse(BASE_REDIRECT_URI))
query = dict(parse_qsl(url_parts[4]))

# Add your token
query.update({"token": API_TOKEN})
url_parts[4] = urlencode(query)

REDIRECT_URI = urlunparse(url_parts)

def _start_gmail_oauth(email_hint: str):
    app.logger.info("Enter _start_gmail_oauth function")
    """Build Google OAuth flow for Gmail with offline access."""
    flow = Flow.from_client_config(
        {
            "web": {
                "client_id": os.getenv("GOOGLE_CLIENT_ID"),
                "client_secret": os.getenv("GOOGLE_CLIENT_SECRET"),
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
                "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs"
            }
        },
        scopes=SCOPES,
        redirect_uri=REDIRECT_URI
    )
    # Request offline access (refresh token) and re-show consent if needed
    auth_url, _ = flow.authorization_url(
        access_type='offline',
        prompt='consent',
        login_hint=email_hint
    )
    app.logger.info("Exit _start_gmail_oauth function")
    return auth_url

@app.route('/signup-jobseeker', methods=['POST'])
def signup_jobseeker():
    app.logger.info("Enter signup_jobseeker function")
    try:
        data = request.form

        # Duplicate checks
        if User.query.filter_by(email=data.get('email')).first():
            return jsonify({'error': 'User already exists'}), 400
        if Candidate.query.filter_by(email=data.get('email')).first():
            return jsonify({'error': 'Candidate already exists'}), 400

        # Parse fields
        first_name = (data.get('first_name') or '').strip()
        last_name  = (data.get('last_name')  or '').strip()
        email      = (data.get('email')      or '').strip()
        phone      = (data.get('phone')      or '').strip()
        password   = (data.get('password')   or '').strip()
        if not all([first_name, last_name, email, password]):
            return jsonify({'error': 'Missing required fields'}), 400

        # Create user
        new_user = User(
            first_name=first_name,
            last_name=last_name,
            email=email,
            phone=phone,
            password=generate_password_hash(password),
            role='jobseeker',
            is_admin=False,
            is_approved=True,
            company_id=None
        )
        db.session.add(new_user)
        db.session.flush()  # get new_user.id

        # Create candidate (self-owned)
        candidate = Candidate(
            recruiter_id=new_user.id,
            name=f"{first_name} {last_name}".strip(),
            email=email,
            phone=phone,
            password=generate_password_hash(password),
            resume_path=None
        )
        db.session.add(candidate)
        db.session.commit()

        # Auto-login + keep candidate_id for oauth2callback
        session['user_id'] = new_user.id
        session['role'] = 'jobseeker'
        session['candidate_id'] = candidate.id

        # Immediately kick off Gmail OAuth consent
        try:
            auth_url = _start_gmail_oauth(email_hint=email)
            return jsonify({'message': 'Signup successful', 'redirect': auth_url})
        except Exception as oe:
            app.logger.error(f"[OAuth Init Error] {oe}")
            # Fallback: go to dashboard without Gmail linking
            return jsonify({'message': 'Signup successful', 'redirect': url_for('login_success')})

    except Exception as e:
        db.session.rollback()
        app.logger.error("Error in signup_jobseeker: %s", traceback.format_exc())
        return jsonify({'error': str(e)}), 500


@app.route('/approve-recruiter/<int:user_id>')
def approve_recruiter(user_id):
    app.logger.info("Enter approve_recruiter function")
    user = User.query.get(user_id)
    if not user:
        return "Recruiter not found", 404
    user.is_approved = True
    db.session.commit()
    return f"Recruiter {user.email} approved! They can now login."


@app.route('/login-success')
def login_success():
    app.logger.info("Enter login_success function")
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login', role='recruiter'))
    user = User.query.get(user_id)
    if not user:
        session.clear()
        return redirect(url_for('login', role='recruiter'))

    if user.role == 'admin':
        return render_template('recruiter_dashboard.html', company=user.company, api_token=os.getenv('API_TOKEN'),backend_url=os.getenv('BACKEND_URL'))
    
    elif user.role == 'jobseeker':
            candidate = Candidate.query.filter_by(email=user.email).first()
            if candidate:
                # ✅ Store candidate_id in session
                session['candidate_id'] = candidate.id
                applied_count = Job.query.filter(
                    Job.candidate_id == candidate.id,
                    Job.status.ilike("applied%")
                ).count()
                return render_template(
                    'job_seeker_dashboard_mockup.html',
                    candidate_id=candidate.id,
                    full_name=candidate.name,
                    email=candidate.email,
                    jobs_applied=applied_count,
                    api_token=os.getenv('API_TOKEN'),
                    backend_url=os.getenv('BACKEND_URL')
                )
            else:
                return "Candidate record not found", 404

    elif user.role == 'recruiter':
        return redirect(url_for('dashboard'))

    return redirect(url_for('login', role='recruiter'))


@app.route('/recruiter-dashboard/<int:recruiter_id>')
def recruiter_dashboard(recruiter_id):
    app.logger.info(f"Enter recruiter_dashboard function recruiter_id : {recruiter_id}")
    if 'user_id' not in session:
        return redirect(url_for('login'))

    current_user = User.query.get(session['user_id'])
    app.logger.info(f"Inside recruiter_dashboard function current_user : {current_user}")
    recruiter = User.query.get(recruiter_id)
    app.logger.info(f"Inside recruiter_dashboard function current_user.role : {current_user.role}")
    # ✅ If current user is not admin or recruiter mismatch — block
    if current_user.role != 'admin' or recruiter.company_id != current_user.company_id:
        return "Unauthorized", 403

    # ✅ Only show candidates of that recruiter
    candidates = Candidate.query.filter_by(recruiter_id=recruiter.id).all()
    app.logger.info(f"Inside recruiter_dashboard function candidates : {candidates}")

    return render_template('dashboard.html', recruiter=recruiter, candidates=candidates, is_admin_view=True, api_token=os.getenv('API_TOKEN'),backend_url=os.getenv('BACKEND_URL'))


@app.route('/admin-dashboard')
def admin_dashboard():
    app.logger.info("Enter admin_dashboard function")
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    if not user or not user.is_admin:
        return "Unauthorized", 403

    company = user.company
    recruiters = User.query.filter_by(company_id=company.id, role='recruiter').all()

    return render_template('recruiter_dashboard.html', company=company, recruiters=recruiters, api_token=os.getenv('API_TOKEN'),backend_url=os.getenv('BACKEND_URL'))


@app.route('/api/recruiters', methods=['GET'])
def get_recruiters():
    app.logger.info("Enter get_recruiters function")
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    user = User.query.get(session['user_id'])
    recruiters = User.query.filter_by(company_id=user.company_id, role='recruiter').all()
    data = [{'id': r.id, 'name': f"{r.first_name} {r.last_name}".strip(), 'email': r.email} for r in recruiters]
    return jsonify({'recruiters': data})

@app.route('/api/recruiters', methods=['POST'])
def api_add_recruiter():
    app.logger.info("Enter api_add_recruiter function")
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    user = User.query.get(session['user_id'])
    app.logger.info(f"Inside api_add_recruiter function user : {user}")
    data = request.get_json()
    name = data['name']
    app.logger.info(f"Inside api_add_recruiter function name : {name}")
    email = data['email']
    app.logger.info(f"Inside api_add_recruiter function email : {email}")
    password = generate_password_hash(data['password'])

    existing_recruiters = User.query.filter_by(company_id=user.company_id, role='recruiter').count()
    if existing_recruiters >= 10:
        return jsonify({'success': False, 'message': 'Maximum 10 recruiters allowed per company'}), 400

    if User.query.filter_by(email=email).first():
        return jsonify({'success': False, 'message': 'Email already exists'}), 400


    new_recruiter = User(
        first_name=name.split()[0],
        last_name=name.split()[1] if ' ' in name else '',
        email=email,
        password=password,
        role='recruiter',
        company_id=user.company_id,
        is_approved=True
    )
    db.session.add(new_recruiter)
    db.session.commit()
    app.logger.info(f"Inside api_add_recruiter function new_recruiter : {new_recruiter}")
    return jsonify({'success': True, 'message': 'Recruiter added successfully'})

@app.route('/api/recruiters/<int:recruiter_id>', methods=['DELETE'])
def delete_recruiter(recruiter_id):
    app.logger.info("Enter delete_recruiter function")
    recruiter = User.query.get(recruiter_id)
    if not recruiter:
        return jsonify({'success': False, 'message': 'Recruiter not found'}), 404
    db.session.delete(recruiter)
    db.session.commit()
    return jsonify({'success': True, 'message': 'Recruiter deleted'})


@app.route('/add-recruiter', methods=['POST'])
def add_recruiter():
    app.logger.info("Enter add_recruiter function")
    user = User.query.get(session['user_id'])
    if user.role != 'admin':
        return "Only admin can add recruiters", 403

    name = request.form['name']
    email = request.form['email']
    password = generate_password_hash(request.form['password'])

    new_recruiter = User(first_name=name, last_name='', email=email, password=password, phone='', role='recruiter', company_id=user.company_id)

    db.session.add(new_recruiter)
    db.session.commit()
    return jsonify({'message': 'Recruiter added'})

@app.route('/api/candidates', methods=['GET'])
def api_get_candidates():
    app.logger.info("Enter api_get_candidates function")
    user_id = session.get('user_id')
    app.logger.info(f"Enter job_dashboard function user_id : {user_id}")
    
    if not user_id:
        return jsonify({'error': 'Unauthorized'}), 401

    user = User.query.get(user_id)
    app.logger.info(f"Inside api_get_candidates function user :  {user}")
    
    if not user:
        return jsonify({'error': 'Invalid session'}), 403

    recruiter_id = request.args.get('recruiter_id', type=int)
    app.logger.info(f"Inside api_get_candidates function recruiter_id :  {recruiter_id}")
    app.logger.info(f"Inside api_get_candidates function user.role :  {user.role}")
    
    if user.role == 'recruiter':
        candidates = Candidate.query.filter_by(recruiter_id=user.id).all()
        app.logger.info(f"Inside api_get_candidates function candidates :  {candidates}")
    elif user.role == 'admin':
        if recruiter_id:
            # Admin fetching a specific recruiter’s candidates
            recruiter = User.query.get(recruiter_id)
            app.logger.info(f"Inside api_get_candidates function recruiter :  {recruiter}")
            if not recruiter or recruiter.company_id != user.company_id:
                return jsonify({'error': 'Unauthorized recruiter ID'}), 403
            candidates = Candidate.query.filter_by(recruiter_id=recruiter_id).all()
            app.logger.info(f"Inside api_get_candidates function candidates :  {candidates}")
        else:
            # Admin fetching all company recruiters’ candidates
            recruiters = User.query.filter_by(company_id=user.company_id, role='recruiter').all()
            app.logger.info(f"Inside api_get_candidates function recruiters :  {recruiters}")
            recruiter_ids = [r.id for r in recruiters]
            candidates = Candidate.query.filter(Candidate.recruiter_id.in_(recruiter_ids)).all()
            app.logger.info(f"Inside api_get_candidates else function candidates :  {candidates}")  
    else:
        return jsonify({'error': 'Invalid user role'}), 403

    data = [{
        'id': c.id,
        'name': c.name,
        'email': c.email,
        'resume_path': c.resume_path,
        'verified': True
    } for c in candidates]

    return jsonify({'candidates': data})


@app.route('/dashboard')
def dashboard():
    app.logger.info("Enter dashboard function")
    user = User.query.get(session['user_id'])
    if user.role != 'recruiter':
        return "Unauthorized", 403

    candidates = Candidate.query.filter_by(recruiter_id=user.id).all()
    return render_template('dashboard.html', recruiter=user, candidates=candidates, api_token=os.getenv('API_TOKEN'),backend_url=os.getenv('BACKEND_URL'))


@app.route('/add-candidate', methods=['POST'])
def add_candidate():
    app.logger.info("Enter add_candidate function")
    current_user = User.query.get(session.get('user_id'))
    app.logger.info(f"Inside add_candidate function current_user : {current_user}")
    if not current_user:
        return jsonify({'error': 'Unauthorized'}), 403

    name = request.form.get('name')
    app.logger.info(f"Inside add_candidate function name : {name}")
    email = request.form.get('email')
    app.logger.info(f"Inside add_candidate function email : {email}")
    resume = request.files.get('resume')
    recruiter_id = request.form.get('recruiter_id')

    if not all([name, email, resume]):
        return jsonify({'error': 'Missing required fields'}), 400

    # Find the recruiter the candidate belongs to
    recruiter = User.query.get(int(recruiter_id))
    if not recruiter:
        return jsonify({'error': 'Recruiter not found'}), 404

    s3_path = save_resume_to_apify(resume.stream, resume.filename)

    candidate = Candidate(
        recruiter_id=recruiter.id,
        name=name,
        email=email,
        resume_path=s3_path
    )
    db.session.add(candidate)
    db.session.commit()

    session['candidate_email'] = email
    session['candidate_id'] = candidate.id

    app.logger.info(f"Candidate {candidate.id} created for recruiter {recruiter.id}")
    # If OAuth flow not required for admin, return normal success
    app.logger.info(f"Inside add_candidate function current_user.role : {current_user.role}")
    app.logger.info(f"Inside add_candidate function recruiter.role : {recruiter.role}")
    if recruiter.role == 'admin':
        return jsonify({'message': 'Candidate added successfully'})
    
    # For recruiters, start Gmail OAuth
    try:
        # The Apify token is now passed in the state parameter within _start_gmail_oauth
        auth_url = _start_gmail_oauth(email_hint=email) 
        api_token = os.getenv('API_TOKEN')
        # The redirect URL must include the Apify token
        return jsonify({'message': 'Candidate added, redirecting to OAuth', 'redirect': auth_url})
    except Exception as e:
        print(f"[ERROR] OAuth start error: {e}")
        app.logger.error(f"[ERROR] OAuth start error: {e}")
        return jsonify({'error': f"OAuth start error: {e}"}), 500

@app.route('/oauth2callback')
def oauth2callback():
    app.logger.info("Enter oauth2callback function")
    try:
        flow = Flow.from_client_config(
            {
                "web": {
                    "client_id": os.getenv("GOOGLE_CLIENT_ID"),
                    "client_secret": os.getenv("GOOGLE_CLIENT_SECRET"),
                    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                    "token_uri": "https://oauth2.googleapis.com/token",
                    "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs"
                }
            },
            scopes=SCOPES,
            redirect_uri=REDIRECT_URI
        )
        flow.fetch_token(authorization_response=request.url)
        creds = flow.credentials

        candidate_id = session.get('candidate_id')
        if not candidate_id:
            return "Missing candidate session", 400

        candidate = Candidate.query.get(candidate_id)
        if not candidate:
            return "Candidate not found", 404

        # Save tokens to the candidate's DB record
        candidate.oauth_token = creds.token
        candidate.refresh_token = creds.refresh_token
        candidate.token_uri = creds.token_uri
        candidate.client_id = creds.client_id
        candidate.client_secret = creds.client_secret
        candidate.scopes = ",".join(creds.scopes)
        db.session.commit()

        # ✅ Add the API token to the redirect URL
        api_token = os.getenv('API_TOKEN')
        dashboard_url = url_for('job_dashboard', candidate_id=candidate.id)
        
        return redirect(f"{dashboard_url}?token={api_token}")
    except Exception as e:
        print(f"[ERROR] OAuth callback error: {e}")
        app.logger.error(f"[ERROR] OAuth callback error: {e}")
        return jsonify({'error': f"OAuth callback error: {e}"}), 500

@app.route('/job-dashboard/<int:candidate_id>')
def job_dashboard(candidate_id):
    app.logger.info(f"Enter job_dashboard function : {candidate_id}")
    candidate = Candidate.query.get(candidate_id)
    if not candidate:
        return "Candidate not found", 404

    # Fetch jobs associated with this candidate
    jobs = Job.query.filter_by(candidate_id=candidate_id).order_by(Job.created_at.desc()).all()

    # Count jobs with status 'Applied' (or adjust as needed)
    applied_count = sum(1 for job in jobs if (job.status or "").lower() == "applied")

    return render_template('job_seeker_dashboard_mockup.html', full_name=candidate.name, candidate_id=candidate.id, email=candidate.email, resume_path=candidate.resume_path, jobs=jobs, applied_count=applied_count, api_token=os.getenv('API_TOKEN'),backend_url=os.getenv('BACKEND_URL'))

@app.route('/activate-job-jarvis/<int:candidate_id>', methods=['POST'])
def activate_job_jarvis(candidate_id):
    app.logger.info(f"Enter activate_job_jarvis function : {candidate_id}")
    from linkedin_bot_fixed import run_linkedin_for_candidate

    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'User not logged in'}), 401

    candidate = Candidate.query.get(candidate_id)
    current_user = User.query.get(session['user_id'])

    if not candidate:
        return jsonify({'success': False, 'message': 'Candidate not found'}), 404

    # Allow recruiter who owns candidate OR admin from same company
    if not (
        candidate.recruiter_id == current_user.id or
        (current_user.role == 'admin' and current_user.company_id == User.query.get(candidate.recruiter_id).company_id)
    ):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    keywords = (candidate.key_skills or "").split(",")
    location = candidate.preferred_location or "remote"
    remote = True

    if not any(keywords):
        keywords = ["software", "developer"]

    # ✅ Scheduler is now managed by scheduler_runner.py
    # No need to call start_bg_scheduler() or schedule_day(candidate_id) here

    def run_scraper_only():
        app.logger.info("Enter run_scraper_only function")
        with app.app_context():
            candidate_id = candidate.id
            try:
                app.logger.info(f"Starting initial scraper for {keywords=} {location=}")
                jobs = main(
                    keywords=keywords,
                    location=location,
                    remote=remote,
                    user_id=candidate.id,
                    db=db,
                    Job=Job,
                    title=candidate.title or ""
                )
                for job in jobs:
                    link = job.get("link", "").lower()
                    if "linkedin" in link:
                        threading.Thread(target=run_linkedin_for_candidate, args=(candidate_id, app)).start()
                    elif "indeed" in link:
                        threading.Thread(target=run_indeed_for_candidate, args=(candidate_id, app)).start()

            except Exception as e:
                app.logger.exception(f"Initial scrape failed for candidate {candidate_id}: {e}")

    threading.Thread(target=run_scraper_only, daemon=True).start()

    # ✅ return a valid response for the frontend fetch()
    return jsonify({
        'success': True,
        'message': 'JobJarvis activated. Initial scrape started.'
    })

@app.route('/pause-job-jarvis/<int:candidate_id>', methods=['POST'])
def pause_job_jarvis(candidate_id):
    app.logger.info(f"Enter pause_job_jarvis function : {candidate_id}")
    # Optional: Set pause flag in DB if implementing real pause logic
    return jsonify({'success': True, 'message': 'JobJarvis paused'})

@app.route('/profile/<int:candidate_id>')
def profile_section(candidate_id):
    app.logger.info(f"Enter profile_section function : {candidate_id}")
    if 'user_id' not in session:
        return redirect('/login')

    candidate = Candidate.query.get(candidate_id)
    current_user = User.query.get(session['user_id'])

    if not candidate:
        return "Candidate not found", 404

    # ✅ Jobseeker can only see their own profile
    if current_user.role == 'jobseeker':
        my_candidate = Candidate.query.filter_by(email=current_user.email).first()
        if my_candidate and my_candidate.id == candidate_id:
            return render_template(
                'profile_section.html',
                candidate_id=candidate.id,
                full_name=candidate.name,
                email=candidate.email,
                phone=candidate.phone,
                title=candidate.title,
                resume_path=candidate.resume_path,
                preferred_location=candidate.preferred_location,
                state=candidate.state, 
                work_authorization=candidate.work_authorization,
                key_skills=candidate.key_skills,
                employment=candidate.employment,
                education=candidate.education,
                certifications=candidate.certifications,
                projects=candidate.projects,
                employment_type=candidate.employment_type,          
                job_type=candidate.job_type,                        
                salary=candidate.salary_expectations,
                remote_only=candidate.remote_only, 
                api_token=os.getenv('API_TOKEN'),
                backend_url=os.getenv('BACKEND_URL')
            )
        else:
            return "Unauthorized access", 403

    # ✅ Recruiter who owns candidate OR admin from same company
    if (
        candidate.recruiter_id == current_user.id or
        (current_user.role == 'admin' and current_user.company_id == User.query.get(candidate.recruiter_id).company_id)
    ):
        return render_template(
            'profile_section.html',
            candidate_id=candidate.id,
            full_name=candidate.name,
            email=candidate.email,
            phone=candidate.phone,
            title=candidate.title,
            resume_path=candidate.resume_path,
            preferred_location=candidate.preferred_location,
            work_authorization=candidate.work_authorization,
            key_skills=candidate.key_skills,
            employment=candidate.employment,
            education=candidate.education,
            certifications=candidate.certifications,
            projects=candidate.projects,
            employment_type=candidate.employment_type,          # NEW
            job_type=candidate.job_type,                        # NEW
            salary=candidate.salary_expectations,               # NEW (template id is "salary")
            remote_only=candidate.remote_only,                  # NEW
            state=candidate.state,                              # NEW
            api_token=os.getenv('API_TOKEN'),
            backend_url=os.getenv('BACKEND_URL')
        )

    return "Unauthorized access", 403

@app.route('/profile', methods=['POST'])
def update_profile():
    candidate_id = request.form.get('candidate_id') or session.get('candidate_id')
    app.logger.info("Enter update_profile function")
    if not candidate_id:
        return jsonify({'error': 'Candidate not logged in'}), 401

    candidate = Candidate.query.get(int(candidate_id))
    app.logger.info(f"Inside update_profile function : {candidate_id}")
    if not candidate:
        return jsonify({'error': 'Candidate not found'}), 404

    # Work authorization update
    if 'work_authorization' in request.form:
        candidate.work_authorization = request.form['work_authorization']

    # ✅ Add employment type update
    if 'employment_type' in request.form:
        candidate.employment_type = request.form['employment_type']

    # Location update
    if 'preferred_location' in request.form:
        candidate.preferred_location = request.form['preferred_location']

    if 'state' in request.form:
        candidate.state = request.form['state']

    current_user = User.query.get(session['user_id'])
    if not current_user:
        return jsonify({'error': 'User not logged in'}), 401

    # ✅ Jobseeker: can only update their own candidate record
    if current_user.role == 'jobseeker':
        my_candidate = Candidate.query.filter_by(email=current_user.email).first()
        if not my_candidate or my_candidate.id != int(candidate_id):
            return jsonify({'error': 'Unauthorized'}), 403

    # ✅ Recruiter/Admin: can update if they own candidate or are allowed as admin
    elif not (
        candidate.recruiter_id == current_user.id or
        (current_user.role == 'admin' and current_user.company_id == User.query.get(candidate.recruiter_id).company_id)
    ):
        return jsonify({'error': 'Unauthorized'}), 403

    action = request.form.get('action')

    # ✅ Handle resume removal
    if action == 'remove_resume':
        candidate.resume_path = None
        db.session.commit()
        return jsonify({'message': 'Resume removed successfully'})

    # ✅ Handle autofill from resume
    elif action == 'autofill_resume':
        import tempfile, os

        if 'resume' in request.files:
            file = request.files['resume']
            if file.filename != '':
                s3_path = save_resume_to_apify(file.stream, secure_filename(file.filename))
                candidate.resume_path = s3_path
                db.session.commit()

        if not candidate.resume_path:
            return jsonify({'error': 'No resume uploaded'}), 400

        # If the resume lives in Spaces, pull a temp copy before parsing
        local_path = candidate.resume_path
        if candidate.resume_path.startswith("apify://"):
            ext = os.path.splitext(candidate.resume_path)[1] or ".pdf"
            tmpf = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
            tmpf.close()
            get_resume_from_apify(candidate.resume_path, tmpf.name)
            local_path = tmpf.name

        parsed = parse_resume(local_path)

        # (Optional) clean up temp file
        try:
            if local_path != candidate.resume_path and os.path.exists(local_path):
                os.unlink(local_path)
        except Exception:
            pass


        # Normalize to the career-info fields we care about
        if isinstance(parsed, dict):
            fields = {
                'key_skills': normalize_field(parsed.get('key_skills'))[:2000],
        		'employment': normalize_field(parsed.get('employment')),
        		'education': normalize_field(parsed.get('education')),
        		'certifications': normalize_field(parsed.get('certifications')),
        		'projects': normalize_field(parsed.get('projects'))
            }
        else:
            # Fallback if parser returns raw text (string)
            text = parsed or ''
            fields = {
                'key_skills': text[:2000],
                'employment': '',
                'education': '',
                'certifications': '',
                'projects': ''
            }

        # Save ONLY career info into DB first
        candidate.key_skills     = fields['key_skills']
        candidate.employment     = fields['employment']
        candidate.education      = fields['education']
        candidate.certifications = fields['certifications']
        candidate.projects       = fields['projects']
        try:
            db.session.commit()
        except SQLAlchemyError as e:
            db.session.rollback()
            app.logger.error("Commit failed in update_profile: %s", traceback.format_exc())
            return jsonify({'error': str(e.__cause__ or e)}), 500

        # Then return the saved values for the frontend to display
        return jsonify({
            'message': 'Resume parsed and saved',
            'profile_data': fields
        })

    # ✅ Handle resume upload only (no action passed)
    if 'resume' in request.files:
        file = request.files['resume']
        if file.filename != '':
            s3_path = save_resume_to_apify(file.stream, secure_filename(file.filename))
            candidate.resume_path = s3_path
            db.session.commit()
            return jsonify({'message': 'Resume uploaded successfully', 'path': s3_path})

    # ✅ Handle profile section updates
    fields = [
        'first_name', 'last_name', 'email', 'phone', 'title',
        'work_authorization', 'preferred_location', 'key_skills',
        'employment', 'education', 'certifications', 'projects',
        'employment_type','job_type','salary_expectations','state'
    ]

    for field in fields:
        if field in request.form:
            value = request.form.get(field)
            if field in ['first_name', 'last_name']:
                name_parts = candidate.name.split()
                if field == 'first_name':
                    candidate.name = f"{value} {name_parts[1]}" if len(name_parts) > 1 else value
                else:
                    candidate.name = f"{name_parts[0]} {value}" if name_parts else value
            else:
                setattr(candidate, field, value)

    # handle checkbox (remote_only)
    remote_only_raw = request.form.get('remote_only')
    if remote_only_raw is not None:
        candidate.remote_only = remote_only_raw in ('1','true','on','yes')

    db.session.commit()
    return jsonify({'message': 'Profile updated successfully'})



# @app.route('/profile', methods=['GET', 'POST'])
# def profile():
#     user = User.query.get(session['user_id'])

#     if request.method == 'POST':
#         if 'resume' in request.files:
#             file = request.files['resume']
#             filename = secure_filename(file.filename)
#             save_path = os.path.join("resumes", filename)
#             file.save(save_path)
#             user.resume_path = save_path
#             db.session.commit()
#             return jsonify({'message': 'Resume uploaded'})
        
#         elif request.form.get('action') == 'remove_resume':
#             candidate_id = request.form.get('candidate_id')
#             candidate = Candidate.query.get(int(candidate_id))
#             candidate.resume_path = None
#             db.session.commit()
#             return jsonify({'message': 'Resume removed'})


#         elif request.form.get('action') == 'autofill_resume':
#             candidate_id = request.form.get('candidate_id')
#             candidate = Candidate.query.get(int(candidate_id))
#             parsed_data = parse_resume(candidate.resume_path)
#             # Optionally map parsed_data to fields
#             return jsonify({'message': 'Autofill success', 'profile_data': parsed_data})


@app.route('/autofill-resume/<int:candidate_id>', methods=['POST'])
def autofill_resume(candidate_id):
    app.logger.info(f"Enter autofill_resume function : {candidate_id}")
    candidate = Candidate.query.get(candidate_id)
    if not candidate or candidate.recruiter_id != session['user_id']:
        return jsonify({'message': 'Unauthorized'}), 403
    if not candidate.resume_path:
        return jsonify({'message': 'No resume uploaded'}), 400

    parsed_text = parse_resume(candidate.resume_path)
    candidate.key_skills = parsed_text[:500]  # You can later improve this logic
    db.session.commit()
    return jsonify({'message': 'Resume data extracted and saved'})

@app.route('/get-applications')
def get_applications():
    candidate_id = request.args.get('candidate_id') or session.get('candidate_id')
    app.logger.info(f"Enter get_applications function : {candidate_id}")
    if not candidate_id:
        return jsonify({'error': 'Candidate not logged in'}), 401

    candidate_id = int(candidate_id)
    jobs = Job.query.filter_by(candidate_id=candidate_id).order_by(Job.created_at.desc()).all()


    applications = []
    for job in jobs:
        applications.append({
            'title': job.title,
            'location': job.location,
            'company': job.company,
            'created_at': job.created_at.strftime('%Y-%m-%d %H:%M') if job.created_at else '',
            'status': job.status or 'Applied',  # Default if missing
            'comment': job.comment or '',
            'interview_status': job.interview_status or '',
            'final_verdict': job.final_verdict or ''
        })

    return jsonify({'status': 'success', 'applications': applications})

@app.route('/get-jobs')
def get_jobs():
    candidate_id = request.args.get('candidate_id') or session.get('candidate_id')
    app.logger.info(f"Enter get_jobs function : {candidate_id}")
    if not candidate_id:
        return jsonify({'status': 'error', 'message': 'Candidate not logged in'}), 401

    candidate_id = int(candidate_id)
    app.logger.info(f"Inside get_jobs function candidate_id : {candidate_id}")
    jobs = Job.query.filter_by(candidate_id=candidate_id).order_by(Job.created_at.desc()).all()
    app.logger.info(f"Inside get_jobs function jobs : {jobs}")

    jobs_data = []
    for job in jobs:
        jobs_data.append({
            'title': job.title,
            'company': job.company,
            'location': job.location,
            'link': job.link,
            'source': job.source,
            'created_at': job.created_at.strftime('%Y-%m-%d %H:%M') if job.created_at else '',
            'status': job.status,
            'comment': job.comment,
            'interview_status': job.interview_status,
            'final_verdict': job.final_verdict
        })

    return jsonify({'status': 'success', 'jobs': jobs_data})


@app.route("/process-jobs", methods=["POST"])
def process_jobs():
    app.logger.info("Enter process_jobs")
    from linkedin_bot_fixed import run_linkedin_for_candidate
    candidate_id = request.form.get("candidate_id") or session.get("candidate_id")
    if not candidate_id:
        return jsonify({"error": "Candidate not logged in"}), 401

    candidate_id = int(candidate_id)
    app.logger.info(f"Inside process_jobs function : {candidate_id}")

    # Get all jobs for the candidate
    jobs = Job.query.filter_by(candidate_id=candidate_id).order_by(Job.created_at.desc()).all()
    if not jobs:
        return jsonify({"error": "No jobs found for candidate"}), 404

    for job in jobs:
        if not job.link:
            continue

        link_lower = job.link.lower()
        if "indeed" in link_lower:
            threading.Thread(target=run_indeed_for_candidate, args=(candidate_id, app)).start()
        elif "linkedin" in link_lower:
            threading.Thread(target=run_linkedin_for_candidate, args=(candidate_id, app)).start()
        else:
            app.logger.warning(f"Unknown source for candidate {candidate_id}: {job.link}")

    return jsonify({"status": "processing_started"})

if __name__ == '__main__':
    app.run(host="0.0.0.0", port=4321)

# ===== Integrated Schedular Code Start =====
# --- Single background scheduler (shared) ---
scheduler = BackgroundScheduler(timezone="UTC")
_scheduler_started = False

def start_bg_scheduler():
    app.logger.info("Enter start_bg_scheduler function")
    global _scheduler_started
    if not _scheduler_started:
        scheduler.start()
        _scheduler_started = True


# --- Helpers for counts / picking jobs / routing runners ---
def monthly_applied_count(candidate_id: int) -> int:
    app.logger.info("Enter monthly_applied_count function")
    # Count only this calendar month
    start = datetime.utcnow().replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    return Job.query.filter(
        Job.candidate_id == candidate_id,
        Job.status.ilike('applied%'),
        Job.created_at >= start
    ).count()

def pick_next_job(candidate_id: int):
    app.logger.info("Enter pick_next_job function")
    # Prefer 'retry' (oldest first), else 'queued' (oldest first)
    retry = Job.query.filter_by(candidate_id=candidate_id, status='retry') \
                     .order_by(Job.created_at.asc()).first()
    if retry:
        return retry
    return Job.query.filter_by(candidate_id=candidate_id, status='queued') \
                    .order_by(Job.created_at.asc()).first()

def apply_one_job(candidate_id: int):
    app.logger.info(f"Enter apply_one_job function: {candidate_id}")
    from linkedin_bot_fixed import run_linkedin_for_candidate

    try:
        with app.app_context():
            # Check application limit
            try:
                if monthly_applied_count(candidate_id) >= 1000:
                    app.logger.info(f"[Pipeline] Target reached for candidate {candidate_id}.")
                    return
            except Exception as e:
                app.logger.exception(f"Error checking monthly_applied_count for candidate {candidate_id}: {e}")
                return  # bail out early if count check fails

            # Pick next job
            try:
                job = pick_next_job(candidate_id)
            except Exception as e:
                app.logger.exception(f"Error picking next job for candidate {candidate_id}: {e}")
                return

            if not job:
                app.logger.info(f"[Pipeline] No queued/retry jobs for candidate {candidate_id}.")
                return

            src = (job.source or "").lower()
            link = job.link or ""

            # Route by source/link
            try:
                if "linkedin" in src or "linkedin." in link:
                    run_linkedin_for_candidate(candidate_id, flask_app=app)
                elif "indeed" in src or "indeed." in link:
                    run_indeed_for_candidate(candidate_id, flask_app=app)
                else:
                    # Default: Indeed runner (it handles many ATS links too)
                    run_indeed_for_candidate(candidate_id, flask_app=app)

                app.logger.info(f"apply_one_job finished successfully for candidate {candidate_id}")

            except Exception as e:
                app.logger.exception(f"Error running job apply logic for candidate {candidate_id}: {e}")

    except Exception as outer:
        app.logger.exception(f"apply_one_job crashed unexpectedly for candidate {candidate_id}: {outer}")


def run_daily_scrape(candidate_id: int):
    app.logger.info(f"Enter run_daily_scrape function: {candidate_id}")
    with app.app_context():
        cand = Candidate.query.get(candidate_id)
        if not cand:
            return
        
        MAX_DAILY_JOBS = 100
        MAX_MONTHLY_JOBS = 1000

        if monthly_applied_count(cand.id) >= MAX_MONTHLY_JOBS:
            app.logger.info(f"[DailyScrape] Monthly cap reached for candidate {cand.id}.")
            return
        
        keywords = [s.strip() for s in (cand.key_skills or "").split(",") if s.strip()] or ["software","developer"]
        location = cand.preferred_location or "remote"

        today = datetime.utcnow().date()
        existing = Job.query.filter(
            Job.candidate_id == cand.id,
            Job.created_at >= datetime.combine(today, datetime.min.time()),
            Job.status.in_(['queued', 'retry', 'applied'])
        ).count()

        if existing < MAX_DAILY_JOBS:
            jobs_to_scrape = MAX_DAILY_JOBS - existing
            scraped = main(keywords=keywords, location=location, remote=True, user_id=cand.id, db=db, Job=Job, max_jobs=jobs_to_scrape, title=cand.title or "")
            app.logger.info(f"[DailyScrape] Scraped {len(scraped)} new jobs for candidate {cand.id}")
        else:
            app.logger.info(f"[DailyScrape] Already hit daily cap for candidate {cand.id}")

def schedule_day(candidate_id: int, tz_name="Asia/Kolkata", stagger_index=0):
    app.logger.info(f"Enter schedule_day function: {candidate_id}")
    tz = ZoneInfo(tz_name)

    # Offset in minutes to stagger job start
    stagger_offset = timedelta(minutes=stagger_index * 2)  # 2 min per candidate

    # Daily scrape at 07:30 IST (stagger optional)
    scheduler.add_job(
        func=lambda cid=candidate_id: run_daily_scrape(cid),
        trigger="cron",
        minute="0-59/15",
        hour="8-23,0",
        timezone=tz,
        id=f"daily-scrape-{candidate_id}",
        replace_existing=True,
        next_run_time=datetime.now(tz) + stagger_offset  # stagger start
    )

    # Applications every 15 minutes between 08:00 and 01:40 IST
    scheduler.add_job(
        func=lambda cid=candidate_id: apply_one_job(cid),
        trigger="cron",
        minute="0-59/15",
        hour="8-23,0",
        timezone=tz,
        id=f"apply-interval-{candidate_id}",
        replace_existing=True,
        next_run_time=datetime.now(tz) + stagger_offset  # stagger start
    )

    app.logger.info(f"Scheduled jobs for candidate {candidate_id} with stagger {stagger_offset}")

# ===== Integrated Schedular Code End =====


@app.route('/save-resume-info', methods=['POST'])
def save_resume_info():
    candidate_id = session.get('candidate_id')
    app.logger.info(f"Enter save_resume_info function: {candidate_id}")
    if not candidate_id:
        return jsonify({"error": "Not logged in"}), 401

    candidate = Candidate.query.get(candidate_id)
    if not candidate:
        return jsonify({"error": "Candidate not found"}), 404

    title = request.form.get('title')
    key_skills = request.form.get('key_skills')  # comma-separated string
    print("Received skills:", key_skills)

    updated = False

    if title:
        candidate.title = title
        updated = True

    if key_skills:
        candidate.key_skills = key_skills
        updated = True

    if updated:
        db.session.commit()
        return jsonify({"message": "Resume info saved successfully"})
    else:
        return jsonify({"error": "No fields provided"}), 400

def normalize_field(val):
    """Convert parsed list/None into safe string for DB."""
    app.logger.info("Enter normalize_field function")
    if isinstance(val, list):
        return ", ".join(str(v).strip() for v in val if v)  # join list into one string
    if val is None:
        return ""
    return str(val).strip()

@app.route('/job-dashboard-data')
def job_dashboard_data():
    candidate_id = session.get("candidate_id")  # assume set during login
    app.logger.info(f"Enter job_dashboard_data function: {candidate_id}")
    if not candidate_id:
        return {"error": "Not logged in"}, 401

    applied_count = Job.query.filter(
        Job.candidate_id == candidate_id,
        Job.status.ilike("applied%")
    ).count()

    return {"applied_count": applied_count}

# ------------------- APIFY STORAGE CONFIG -------------------
from apify_client import ApifyClient

APIFY_TOKEN = os.getenv("API_TOKEN", "")
app.logger.info(f"APIFY_TOKEN from environment : {APIFY_TOKEN}")
apify_client = ApifyClient(APIFY_TOKEN)

# Get or create by name, then get the actual client
store_meta = apify_client.key_value_stores().get_or_create(name="resumes")
app.logger.info(f"Get or create by name, then get the actual client. store_meta : {store_meta}")
resumes_store = apify_client.key_value_store(store_meta["id"])
app.logger.info(f"Get or create by name, then get the actual client. resumes_store : {resumes_store}")

def save_resume_to_apify(file_obj, original_filename):
    """Save resume file to Apify Key-Value store and return apify:// path"""
    app.logger.info(f"Enter save_resume_to_apify function. file_obj : {file_obj}")
    app.logger.info(f"Inside save_resume_to_apify function. original_filename : {original_filename}")
    ext = os.path.splitext(original_filename)[1]
    unique_suffix = uuid.uuid4().hex[:6]
    filename = f"{os.path.splitext(original_filename)[0]}_{unique_suffix}{ext}"

    key = f"resume-{filename}"
    file_bytes = file_obj.read()
    resumes_store.set_record(key, file_bytes)   # ✅ uses store by name safely

    app.logger.info(f"Exit save_resume_to_apify function. key : {key}")
    return f"apify://resumes/{key}"

def get_resume_from_apify(apify_path, local_temp_path):
    """Download resume from Apify KV store to a local temp file"""
    app.logger.info(f"Enter get_resume_from_apify function. apify_path : {apify_path}")
    app.logger.info(f"Inside get_resume_from_apify function. local_temp_path : {local_temp_path}")
    key = apify_path.replace("apify://resumes/", "")
    record = resumes_store.get_record(key)

    if record and "value" in record:
        with open(local_temp_path, "wb") as f:
            f.write(record["value"])
        app.logger.info(f"Exit get_resume_from_apify function. local_temp_path : {local_temp_path}")
        return local_temp_path
    else:
        raise FileNotFoundError(f"Resume not found in Apify store: {key}")

# Route to view resume securely
@app.route('/view-resume/<int:candidate_id>')
def view_resume(candidate_id):
    app.logger.info("Enter view_resume function apify.")
    """
    Generates a temporary, secure URL to view a private resume from Apify.
    """
    if 'user_id' not in session:
        return "Unauthorized", 401

    candidate = Candidate.query.get(candidate_id)
    app.logger.info(f"Inside view_resume function apify. candidate : {candidate}")
    if not candidate or not candidate.resume_path:
        return "Resume not found", 404       

    key = candidate.resume_path.replace("apify://resumes/", "")
    app.logger.info(f"Inside view_resume function. key : {key}")

    # ✅ Use store_meta, not resumes_store
    store_id = store_meta["id"]
    url = f"https://api.apify.com/v2/key-value-stores/{store_id}/records/{key}?token={APIFY_TOKEN}"

    app.logger.info(f"Exit view_resume function apify. url : {url}")
    return redirect(url)