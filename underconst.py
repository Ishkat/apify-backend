from flask import Flask, render_template, request, jsonify, redirect, url_for, session
from flask_sqlalchemy import SQLAlchemy
from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
import os
import re
import logging
from flask_cors import CORS
# from universal_job_scraper import main as run_scraper
from werkzeug.security import generate_password_hash, check_password_hash
import pdfplumber
import sys
import secrets
from datetime import timedelta
from datetime import datetime
import threading
from selenium import webdriver
import subprocess
from flask import render_template, session, redirect
from werkzeug.utils import secure_filename
# server.py (changes)
from models import db, User, Candidate, Job, Company  # ✅ Import from models.py
from universal_job_scraper import main, save_jobs_to_db  # ✅ Now safe to import
from flask import Flask, request, jsonify, render_template, redirect, session

from flask_migrate import Migrate
from flask import Flask, session
from flask_session import Session
from flask_mail import Mail, Message
from indeed import run_indeed_for_candidate
from linkedin_bot import run_linkedin_for_candidate

import json
from openai import OpenAI
import pytesseract
from PIL import Image
from docx import Document  # For .docx
import mammoth          # For .doc
# import boto3

OPENAI_API_KEY = ""
# Initialize OpenAI client
client = OpenAI(api_key=OPENAI_API_KEY)

# ---------- OCR for scanned PDFs ----------
def extract_text_with_ocr(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            img = page.to_image(resolution=300).original
            text += pytesseract.image_to_string(Image.fromarray(img)) + "\n"
    return text

# ---------- Clean Word text ----------
def clean_word_text(text):
    text = re.sub(r"^[\s]*[\u2022\-\*\d]+\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    text = re.sub(r"\n?\s*\b\d+\b\s*\n?", "\n", text)
    return text.strip()

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Detect table rows separately
    education_table_rows = []
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                row_text = " | ".join(cells)
                tables_text.append(row_text)
                # Detect possible education rows
                if re.search(r"(B\.?Tech|B\.?E\.?|M\.?Tech|Diploma|High School|XII|College|University)",
                             row_text, re.IGNORECASE):
                    education_table_rows.append(row_text)

    full_text = "\n".join(paragraphs + tables_text)
    return clean_word_text(full_text), education_table_rows

def extract_text_from_doc(doc_path):
    with open(doc_path, "rb") as doc_file:
        result = mammoth.extract_raw_text(doc_file)
    return clean_word_text(result.value)

# ---------- Fallback section extraction ----------
# Add this mapping near the top of your file (or inside the module)
SECTION_VARIANTS = {
    "skills": [
        "skills", "skillset", "technical skills", "skills & tools", "software skills", "software tools"
    ],
    "experience": [
        "work experience", "experience", "professional experience", "employment history", "work history"
    ],
    "education": [
        "education", "educational qualifications", "educational qualification",
        "educational", "academics", "educational qualifications:", "educational qualification:"
    ],
    "certifications": [
        "certifications", "certificates", "licenses"
    ],
    "projects": [
        "projects", "personal projects", "academic projects", "project"
    ]
}

def _find_heading_end(text: str, variants: list) -> int:
    """
    Return index (end of match) of the first heading variant found (anchored at line start),
    or -1 if none.
    """
    for v in variants:
        # Anchor to start of line, allow optional whitespace and optional colon
        pattern = rf'(?im)^[\s]*{re.escape(v)}\b[\s:]*'
        m = re.search(pattern, text)
        if m:
            return m.end()
    return -1

def _find_next_heading_pos(text: str, start: int) -> int:
    """
    Find the earliest heading occurrence after `start`. Return its index or len(text) if none.
    """
    end = len(text)
    for variants in SECTION_VARIANTS.values():
        for v in variants:
            pattern = rf'(?im)^[\s]*{re.escape(v)}\b[\s:]*'
            m = re.search(pattern, text[start:], flags=re.IGNORECASE | re.MULTILINE)
            if m:
                pos = start + m.start()
                if pos < end:
                    end = pos
    return end

def extract_sections(text: str) -> dict:
    t = (text or "").strip()
    sections = {"key_skills": "", "employment": "", "education": "", "certifications": "", "projects": ""}

    # Helper to cut block for a given section using SECTION_VARIANTS anchors
    def cut_section(section_name: str) -> str:
        start = _find_heading_end(t, SECTION_VARIANTS.get(section_name, []))
        if start == -1:
            return ""
        end = _find_next_heading_pos(t, start)
        return t[start:end].strip()

    # Extract raw blocks using anchored headings
    raw_skills = cut_section("skills")
    raw_employment = cut_section("experience")
    raw_education = cut_section("education")
    raw_certifications = cut_section("certifications")
    raw_projects = cut_section("projects")

    # If education block is empty or looks malformed, attempt to extract table-like rows / degree lines
    if (not raw_education) or (len(raw_education) < 6) or re.match(r"^[a-z]{1,10}\s+qualification[:\s]*$", raw_education.strip().lower()):
        # find lines that look like education rows (table rows or degree lines)
        lines = [ln.strip() for ln in t.splitlines() if ln.strip()]
        degree_pattern = re.compile(
            r'\b(B\.?E\.?|BTECH|B\.?Tech|Bachelor|Bachelors|B\.?Sc|M\.?Tech|MTech|M\.?E\.?|MBA|XII|Xth|High School|Diploma)\b',
            flags=re.IGNORECASE
        )
        candidate_lines = []
        for ln in lines:
            # Accept table rows using '|' (from docx table export) or lines containing degree keywords or a 4-digit year
            if "|" in ln and re.search(r'\b(19|20)\d{2}\b', ln):
                candidate_lines.append(ln)
            elif degree_pattern.search(ln):
                candidate_lines.append(ln)
            elif re.search(r'\b(19|20)\d{2}\b', ln) and len(ln.split()) < 12:
                candidate_lines.append(ln)
        if candidate_lines:
            # join a few likely lines
            raw_education = "\n".join(candidate_lines[:10])

    # ----------------- Cleaning functions -----------------
    def clean_skills(raw: str) -> str:
        if not raw:
            return ""
        parts = re.split(r"[,;\n]", raw)
        stopwords = {"and","or","the","in","of","with","through","during","by","for","to","a","an","as"}
        bad_verbs = ["developed","designed","collaborated","created","delivered","implemented","achieved","contributed","adapted","enhanced","optimized","built","managed","worked","using","leveraging"]
        skills = []
        for p in parts:
            s = re.sub(r"\s+", " ", p).strip(" .,-")
            if not s or len(s) > 30:
                continue
            words = s.split()
            if len(words) > 4:
                continue
            if any(bv in s.lower() for bv in bad_verbs):
                continue
            if all(w.lower() in stopwords for w in words):
                continue
            skills.append(s)
        # dedupe preserving order
        seen = set()
        out = []
        for sk in skills:
            lk = sk.lower()
            if lk not in seen:
                seen.add(lk)
                out.append(sk)
        return ", ".join(out[:50])

    def clean_employment(raw: str) -> str:
        if not raw:
            return ""
        lines = [l.strip() for l in raw.splitlines() if l.strip()]
        date_pattern = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)?\s?\d{4}"
        cleaned = []
        for line in lines:
            line_clean = re.sub(r"^[\-\•\d\.\)\s]+", "", line)
            start, end = "", ""
            if "-" in line_clean and re.search(date_pattern, line_clean, flags=re.IGNORECASE):
                parts = re.split(r"\s*-\s*", line_clean)
                if len(parts) >= 2:
                    start = parts[0].strip()
                    end = parts[1].strip()
            comp_desig = re.sub(date_pattern, "", line_clean, flags=re.IGNORECASE).strip(" .,-")
            if comp_desig or start or end:
                pieces = []
                if comp_desig:
                    pieces.append(comp_desig)
                if start or end:
                    pieces.append(f"{start} - {end}".strip(" -"))
                cleaned.append(". ".join(pieces) if len(pieces) > 1 else pieces[0])
        return "\n".join(cleaned)

    def parse_education_block(raw: str) -> str:
        if not raw:
            return ""
        lines = [l.strip() for l in raw.splitlines() if l.strip()]
        cleaned = []
        for line in lines:
            # If it's a table row with pipes, try to map cells to columns
            if "|" in line:
                parts = [p.strip() for p in line.split("|") if p.strip()]
                # Heuristic: if first cell is serial/no, then degree is second, inst third, board fourth, year last
                year = ""
                for part in reversed(parts):
                    m = re.search(r'\b(19|20)\d{2}\b', part)
                    if m:
                        year = m.group(0)
                        break
                degree = ""
                institution = ""
                # find institution by presence of 'college','institute','university','school'
                for p in parts:
                    if re.search(r'\b(college|institute|university|school|academy)\b', p, re.IGNORECASE):
                        institution = p
                        break
                # degree candidates
                for p in parts:
                    if re.search(r'\b(B\.?E\.?|BTECH|B\.?Tech|Bachelor|Bachelors|M\.?Tech|MTech|MBA|Diploma|XII|Xth)\b', p, re.IGNORECASE):
                        degree = p
                        break
                # fallback assignment if empty
                if not degree and len(parts) >= 2:
                    degree = parts[1]
                if not institution and len(parts) >= 3:
                    institution = parts[2]
                # Build string "Institution - Degree. Start - End" (we'll set end_date = year)
                entry = institution if institution else degree
                if degree and institution:
                    entry = f"{institution} - {degree}"
                if year:
                    entry += f". {year} - {year}"
                cleaned.append(entry.strip(" .,-"))
            else:
                # Non-table line: try extracting degree/institution/year with regex
                year_match = re.search(r'\b(19|20)\d{2}\b', line)
                year = year_match.group(0) if year_match else ""
                deg_match = re.search(r"((?:B\.?E\.?|BTECH|B\.?Tech|Bachelor(?:'s|\s)?|M\.?Tech|MTech|MBA|Diploma|XII|Xth|High School)[^,;\n]*)", line, re.IGNORECASE)
                degree = deg_match.group(1).strip() if deg_match else ""
                inst_match = re.search(r'(.+?)(?:,|\-| at | in | for | \| )', line)
                institution = ""
                if 'college' in line.lower() or 'institute' in line.lower() or 'university' in line.lower() or 'school' in line.lower():
                    # fallback: take the longer chunk not matching degree/year
                    parts = [p.strip() for p in re.split(r'[,;\-|]', line) if p.strip()]
                    for p in parts:
                        if degree and degree.lower() in p.lower():
                            continue
                        if re.search(r'\b(19|20)\d{2}\b', p):
                            continue
                        if len(p) > 3 and not re.search(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\b', p, re.IGNORECASE):
                            institution = p
                            break
                if not institution:
                    # fallback to using the remainder text after removing degree and year
                    tmp = line
                    if degree:
                        tmp = tmp.replace(degree, "")
                    if year:
                        tmp = tmp.replace(year, "")
                    tmp = re.sub(r'[\|\-\:,]+', ' ', tmp).strip()
                    if len(tmp) > 3:
                        institution = tmp
                entry = (institution or degree or line).strip()
                if year:
                    entry += f". {year} - {year}"
                cleaned.append(entry.strip(" .,-"))
        # dedupe
        seen = set()
        out = []
        for c in cleaned:
            lk = c.lower()
            if lk not in seen:
                seen.add(lk)
                out.append(c)
        return "\n".join(out)

    # build final cleaned sections
    sections["key_skills"] = clean_skills(raw_skills)
    sections["employment"] = clean_employment(raw_employment)
    sections["education"] = parse_education_block(raw_education)
    sections["certifications"] = (raw_certifications.replace("\n", ", ") or "").strip()
    sections["projects"] = "\n".join([ln.strip() for ln in raw_projects.splitlines() if ln.strip()]) if raw_projects else ""

    # last resort: if all empty, set skills to first chunk so UI doesn't break
    if not any(sections.values()):
        sections["key_skills"] = t[:2000]

    return sections

# ---------- Main parser ----------
def parse_resume(resume_path):
    ext = os.path.splitext(resume_path)[1].lower()
    text = ""
    edu_rows = []  # for docx table education override

    try:
        if ext == ".pdf":
            with pdfplumber.open(resume_path) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            if not text.strip():
                print("[parse_resume] Using OCR for scanned PDF")
                text = extract_text_with_ocr(resume_path)
        elif ext == ".docx":
            text, edu_rows = extract_text_from_docx(resume_path)
        elif ext == ".doc":
            text = extract_text_from_doc(resume_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    except Exception as e:
        print(f"[parse_resume] Error reading file: {e}")
        return {k: "" for k in ["key_skills","employment","education","certifications","projects"]}

    # If OpenAI fails, fallback
    try:
        prompt = f""" ... same as before ... """
        response = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt,
            temperature=0,
            max_output_tokens=1000
        )
        parsed = json.loads(response.output_text)

        # If docx table rows contain education data, use them to replace education field
        # Force education to come from table rows if available
        if edu_rows:
            edu_list = []
            for row in edu_rows:
                parts = [p.strip() for p in row.split("|") if p.strip()]
                degree = ""
                institution = ""
                year = ""

                # Guess columns: [SrNo, Degree, Institution, Board, Year]
                if len(parts) >= 2:
                    degree = parts[1]
                if len(parts) >= 3:
                    institution = parts[2]

                # Extract year from any column
                for p in parts:
                    m = re.search(r"\b(19|20)\d{2}\b", p)
                    if m:
                        year = m.group(0)
                        break

                # Only keep valid-looking entries
                if institution or degree:
                    edu_list.append({
                        "institution": institution,
                        "degree": degree,
                        "start_date": str(int(year) - 4) if year and "b" in degree.lower() else "",
                        "end_date": year
                    })

            parsed["education"] = edu_list

        # Format output for UI
        return {
            "key_skills": ", ".join(parsed.get("key_skills", [])) or "Not specified",
            "employment": "\n".join([
                f"{job.get('company', '')}" +
                (f" - {job.get('designation', '')}" if job.get("designation") else "") +
                (f". {job.get('start_date', '')} - {job.get('end_date', '')}" if job.get("start_date") or job.get("end_date") else "")
                for job in parsed.get("employment", [])
                if job.get("company")
            ]),
            "education": "\n".join([
                f"{edu.get('institution', '')}" +
                (f" - {edu.get('degree', '')}" if edu.get("degree") else "") +
                (f". {edu.get('start_date', '')} - {edu.get('end_date', '')}" if edu.get("start_date") or edu.get("end_date") else "")
                for edu in parsed.get("education", [])
                if edu.get("institution")
            ]),
            "certifications": ", ".join(parsed.get("certifications", [])),
            "projects": "\n".join(parsed.get("projects", []))
        }
    except Exception as e:
        print(f"[parse_resume] OpenAI parsing failed: {e}")
        return extract_sections(text)

# Example usage
if __name__ == "__main__":
    resume_file = "Bhavesh Walankar Resume.pdf"
    parsed_data = parse_resume(resume_file)
    print(json.dumps(parsed_data, indent=2))


app = Flask(__name__, template_folder='templates')
app.secret_key = secrets.token_hex(16)
# Enable CORS
CORS(app, supports_credentials=True)

# Configure MySQL
app.config['SQLALCHEMY_DATABASE_URI'] = ''

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Session configuration
app.config['SESSION_TYPE'] = 'sqlalchemy'
app.config['SESSION_SQLALCHEMY'] = db
app.config['SESSION_COOKIE_SAMESITE'] = "None"      # Required if frontend and backend are on different domains
app.config['SESSION_COOKIE_SECURE'] = True          # Required for HTTPS
app.config['SESSION_COOKIE_HTTPONLY'] = True        # Extra security


app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'dr.nayakinaidu@gmail.com'
app.config['MAIL_PASSWORD'] = ''
app.config['MAIL_DEFAULT_SENDER'] = 'dr.nayakinaidu@gmail.com'

mail = Mail(app)

# Initialize extensions
db.init_app(app)
Session(app)
migrate = Migrate(app, db)  # Add this line

# Create DB tables (including session table)
with app.app_context():
    db.create_all()


# Configure logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
app.logger.setLevel(logging.DEBUG)

# Directory to store resumes
RESUME_DIR = "resumes"
if not os.path.exists(RESUME_DIR):
    os.makedirs(RESUME_DIR)

@app.route('/')
def landing():
    return render_template('landing.html')

@app.route('/login')
def login():
    role = request.args.get('role')
    session['login_role'] = role
    return render_template('login_recruiter.html') if role == 'recruiter' else render_template('login.html', role=role)

@app.route('/login', methods=['POST'])
def login_submit():
    email = request.form.get('email')
    password = request.form.get('password')
    role = session.get('login_role')

    user = User.query.filter_by(email=email).first()
    if not user or not check_password_hash(user.password, password):
        return jsonify({'error': 'Invalid credentials'}), 401
    # Only allow login for approved users

    if not user.is_approved:
        return jsonify({'error': 'Your account is not approved yet.'}), 403
    
    # Ensure role matches
    # Allow jobseekers to log in from employee portal
    if role == 'employee' and user.role == 'jobseeker':
        pass  # allow

    elif user.role != role and not (user.role == 'admin' and role == 'recruiter'):
        return jsonify({'error': f'You are a {user.role}, not a {role}'}), 403
    
    session['user_id'] = user.id
    session['email'] = user.email
    session['role'] = user.role
        # 🔁 This is the part you replace
    # ✅ Fix this block:
    redirect_url = url_for('login_success')

    return jsonify({'message': 'Login successful','role':user.role, 'redirect': redirect_url})

@app.route('/signup', methods=['POST'])
def signup():
    try:
        data = request.form
        if User.query.filter_by(email=data['email']).first():
            return jsonify({'error': 'User already exists'}), 400

        company = Company(name=data['company_name'], location=data['company_location'])
        db.session.add(company)
        db.session.flush()

        new_user = User(
            first_name=data['first_name'],
            last_name=data['last_name'],
            email=data['email'],
            phone=data['phone'],
            password=generate_password_hash(data['password']),
            role='admin',  # ✅ Set admin role explicitly
            company_id=company.id,
            is_approved=False  # ✅ Recruiters require approval
        )
        db.session.add(new_user)
        db.session.flush()  # Ensure new_user.id is available

        # ✅ Send approval email to admin
        approval_link = url_for('approve_recruiter', user_id=new_user.id, _external=True)
        msg = Message(
            subject="New Recruiter Signup Approval Needed",
            recipients=['dr.nayakinaidu@gmail.com'],  # 🔁 Replace this
            body=f"A new recruiter ({new_user.email}) signed up.\n\nApprove them here: {approval_link}"
        )
        mail.send(msg)

        db.session.commit()

        # Don't set session — user should not be logged in yet
        return jsonify({'message': 'Signup successful. Your request is pending approval. You will be notified once approved.'})


    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

from werkzeug.security import generate_password_hash
import traceback

@app.route('/signup-jobseeker', methods=['POST'])
def signup_jobseeker():
    try:
        data = request.form

        if User.query.filter_by(email=data.get('email')).first():
            return jsonify({'error': 'User already exists'}), 400
        if Candidate.query.filter_by(email=data.get('email')).first():
            return jsonify({'error': 'Candidate already exists'}), 400

        first_name = data.get('first_name', '').strip()
        last_name = data.get('last_name', '').strip()
        email = data.get('email', '').strip()
        phone = data.get('phone', '').strip()
        password = data.get('password', '').strip()

        if not all([first_name, last_name, email, password]):
            return jsonify({'error': 'Missing required fields'}), 400

        # Create user
        new_user = User(
            first_name=first_name,
            last_name=last_name,
            email=email,
            phone=phone,
            password=generate_password_hash(password),
            role='jobseeker',  # ✅ matches login flow
            is_admin=False,
            is_approved=True,
            company_id=None
        )
        db.session.add(new_user)
        db.session.flush()

        # Create candidate linked to same user
        candidate = Candidate(
            recruiter_id=new_user.id,  # satisfies FK
            name=f"{first_name} {last_name}",
            email=email,
            phone=phone,
            password=generate_password_hash(password),
            resume_path=None
        )
        db.session.add(candidate)
        db.session.commit()

        # ✅ Auto-login
        session['user_id'] = new_user.id
        session['role'] = 'jobseeker'

        return redirect(url_for('login_success'))

    except Exception as e:
        db.session.rollback()
        app.logger.error("Error in signup_jobseeker: %s", traceback.format_exc())
        return jsonify({'error': str(e)}), 500


@app.route('/approve-recruiter/<int:user_id>')
def approve_recruiter(user_id):
    user = User.query.get(user_id)
    if not user:
        return "Recruiter not found", 404
    user.is_approved = True
    db.session.commit()
    return f"Recruiter {user.email} approved! They can now login."


@app.route('/login-success')
def login_success():
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login', role='recruiter'))
    user = User.query.get(user_id)
    if not user:
        session.clear()
        return redirect(url_for('login', role='recruiter'))

    if user.role == 'admin':
        return render_template('recruiter_dashboard.html', company=user.company)
    
    elif user.role == 'jobseeker':
            candidate = Candidate.query.filter_by(email=user.email).first()
            if candidate:
                # ✅ Store candidate_id in session
                session['candidate_id'] = candidate.id
                return render_template(
                    'job_seeker_dashboard_mockup.html',
                    candidate_id=candidate.id,
                    full_name=candidate.name,
                    email=candidate.email
                )
            else:
                return "Candidate record not found", 404

    elif user.role == 'recruiter':
        return redirect(url_for('dashboard'))

    return redirect(url_for('login', role='recruiter'))


@app.route('/recruiter-dashboard/<int:recruiter_id>')
def recruiter_dashboard(recruiter_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))

    current_user = User.query.get(session['user_id'])
    recruiter = User.query.get(recruiter_id)

    # ✅ If current user is not admin or recruiter mismatch — block
    if current_user.role != 'admin' or recruiter.company_id != current_user.company_id:
        return "Unauthorized", 403

    # ✅ Only show candidates of that recruiter
    candidates = Candidate.query.filter_by(recruiter_id=recruiter.id).all()

    return render_template('dashboard.html', recruiter=recruiter, candidates=candidates, is_admin_view=True)


@app.route('/admin-dashboard')
def admin_dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    if not user or not user.is_admin:
        return "Unauthorized", 403

    company = user.company
    recruiters = User.query.filter_by(company_id=company.id, role='recruiter').all()

    return render_template('recruiter_dashboard.html', company=company, recruiters=recruiters)


@app.route('/api/recruiters', methods=['GET'])
def get_recruiters():
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    user = User.query.get(session['user_id'])
    recruiters = User.query.filter_by(company_id=user.company_id, role='recruiter').all()
    data = [{'id': r.id, 'name': f"{r.first_name} {r.last_name}".strip(), 'email': r.email} for r in recruiters]
    return jsonify({'recruiters': data})

@app.route('/api/recruiters', methods=['POST'])
def api_add_recruiter():
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    user = User.query.get(session['user_id'])
    data = request.get_json()
    name = data['name']
    email = data['email']
    password = generate_password_hash(data['password'])

    existing_recruiters = User.query.filter_by(company_id=user.company_id, role='recruiter').count()
    if existing_recruiters >= 10:
        return jsonify({'success': False, 'message': 'Maximum 10 recruiters allowed per company'}), 400

    if User.query.filter_by(email=email).first():
        return jsonify({'success': False, 'message': 'Email already exists'}), 400


    new_recruiter = User(
        first_name=name.split()[0],
        last_name=name.split()[1] if ' ' in name else '',
        email=email,
        password=password,
        role='recruiter',
        company_id=user.company_id,
        is_approved=True
    )
    db.session.add(new_recruiter)
    db.session.commit()
    return jsonify({'success': True, 'message': 'Recruiter added successfully'})

@app.route('/api/recruiters/<int:recruiter_id>', methods=['DELETE'])
def delete_recruiter(recruiter_id):
    recruiter = User.query.get(recruiter_id)
    if not recruiter:
        return jsonify({'success': False, 'message': 'Recruiter not found'}), 404
    db.session.delete(recruiter)
    db.session.commit()
    return jsonify({'success': True, 'message': 'Recruiter deleted'})


@app.route('/add-recruiter', methods=['POST'])
def add_recruiter():
    user = User.query.get(session['user_id'])
    if user.role != 'admin':
        return "Only admin can add recruiters", 403

    name = request.form['name']
    email = request.form['email']
    password = generate_password_hash(request.form['password'])

    new_recruiter = User(first_name=name, last_name='', email=email, password=password, phone='', role='recruiter', company_id=user.company_id)

    db.session.add(new_recruiter)
    db.session.commit()
    return jsonify({'message': 'Recruiter added'})

@app.route('/api/candidates', methods=['GET'])
def api_get_candidates():
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'error': 'Unauthorized'}), 401

    user = User.query.get(user_id)
    if not user:
        return jsonify({'error': 'Invalid session'}), 403

    recruiter_id = request.args.get('recruiter_id', type=int)

    if user.role == 'recruiter':
        candidates = Candidate.query.filter_by(recruiter_id=user.id).all()

    elif user.role == 'admin':
        if recruiter_id:
            # Admin fetching a specific recruiter’s candidates
            recruiter = User.query.get(recruiter_id)
            if not recruiter or recruiter.company_id != user.company_id:
                return jsonify({'error': 'Unauthorized recruiter ID'}), 403
            candidates = Candidate.query.filter_by(recruiter_id=recruiter_id).all()
        else:
            # Admin fetching all company recruiters’ candidates
            recruiters = User.query.filter_by(company_id=user.company_id, role='recruiter').all()
            recruiter_ids = [r.id for r in recruiters]
            candidates = Candidate.query.filter(Candidate.recruiter_id.in_(recruiter_ids)).all()
    else:
        return jsonify({'error': 'Invalid user role'}), 403

    data = [{
        'id': c.id,
        'name': c.name,
        'email': c.email,
        'resume_path': c.resume_path,
        'verified': True
    } for c in candidates]

    return jsonify({'candidates': data})


@app.route('/dashboard')
def dashboard():
    user = User.query.get(session['user_id'])
    if user.role != 'recruiter':
        return "Unauthorized", 403

    candidates = Candidate.query.filter_by(recruiter_id=user.id).all()
    return render_template('dashboard.html', recruiter=user, candidates=candidates)


SCOPES = [
  'https://www.googleapis.com/auth/gmail.readonly',
  'https://www.googleapis.com/auth/gmail.send'
]

REDIRECT_URI = 'http://127.0.0.1:5002/oauth2callback'

@app.route('/add-candidate', methods=['POST'])
def add_candidate():
    current_user = User.query.get(session.get('user_id'))
    if not current_user:
        return jsonify({'error': 'Unauthorized'}), 403

    # Determine recruiter_id
    if current_user.role == 'recruiter':
        recruiter_id = current_user.id
    elif current_user.role == 'admin':
        recruiter_id = request.form.get('recruiter_id', type=int)
        if not recruiter_id:
            return jsonify({'error': 'Missing recruiter ID'}), 400
        recruiter = User.query.get(recruiter_id)
        if not recruiter or recruiter.company_id != current_user.company_id or recruiter.role != 'recruiter':
            return jsonify({'error': 'Invalid recruiter ID'}), 403
    else:
        return jsonify({'error': 'Invalid role'}), 403

    name = request.form['name']
    email = request.form['email']
    password = request.form.get('password') or request.form.get('candidate_password')
    resume = request.files.get('resume')
    resume_path = None

    if not os.path.exists('resumes'):
        os.makedirs('resumes')
    if resume and resume.filename != '':
        filename = f"{name}_{email.split('@')[0]}.pdf"
        local_resume_path = os.path.join('resumes', filename)
        resume.save(local_resume_path)
        resume_path = local_resume_path

    candidate = Candidate(
        name=name,
        email=email,
        resume_path=resume_path,
        password=generate_password_hash(password) if password else None,
        recruiter_id=recruiter_id
    )
    db.session.add(candidate)
    db.session.commit()

    session['candidate_email'] = email
    session['candidate_id'] = candidate.id

    # If OAuth flow not required for admin, return normal success
    if current_user.role == 'admin':
        return jsonify({'message': 'Candidate added successfully'})
    
    # For recruiters, start Gmail OAuth
    try:
        os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'
        flow = Flow.from_client_secrets_file(
            'credentials.json',
            scopes=SCOPES,
            redirect_uri=REDIRECT_URI
        )
        auth_url, _ = flow.authorization_url(prompt='consent', login_hint=email)
        return jsonify({'message': 'Candidate added, redirecting to OAuth', 'redirect': auth_url})
    except Exception as e:
        print(f"[ERROR] OAuth start error: {e}")
        return jsonify({'error': f"OAuth start error: {e}"}), 500


@app.route('/oauth2callback')
def oauth2callback():
    try:
        flow = Flow.from_client_secrets_file(
            'credentials.json',
            scopes=SCOPES,
            redirect_uri=REDIRECT_URI
        )
        flow.fetch_token(authorization_response=request.url)
        creds = flow.credentials

        candidate_id = session.get('candidate_id')
        if not candidate_id:
            return "Missing candidate session", 400

        candidate = Candidate.query.get(candidate_id)
        if not candidate:
            return "Candidate not found", 404

        # Save tokens to the candidate's DB record
        candidate.oauth_token = creds.token
        candidate.refresh_token = creds.refresh_token
        candidate.token_uri = creds.token_uri
        candidate.client_id = creds.client_id
        candidate.client_secret = creds.client_secret
        candidate.scopes = ",".join(creds.scopes)
        db.session.commit()

        return redirect(url_for('job_dashboard', candidate_id=candidate.id))
    except Exception as e:
        print(f"[ERROR] OAuth callback error: {e}")
        return jsonify({'error': f"OAuth callback error: {e}"}), 500

    
@app.route('/job-dashboard/<int:candidate_id>')
def job_dashboard(candidate_id):
    candidate = Candidate.query.get(candidate_id)
    if not candidate:
        return "Candidate not found", 404

    # Fetch jobs associated with this candidate
    jobs = Job.query.filter_by(candidate_id=candidate_id).order_by(Job.created_at.desc()).all()

    return render_template('job_seeker_dashboard_mockup.html', full_name=candidate.name, candidate_id=candidate.id, email=candidate.email, resume_path=candidate.resume_path, jobs=jobs)

@app.route('/activate-job-jarvis/<int:candidate_id>', methods=['POST'])
def activate_job_jarvis(candidate_id):
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'User not logged in'}), 401

    candidate = Candidate.query.get(candidate_id)
    current_user = User.query.get(session['user_id'])

    if not candidate:
        return jsonify({'success': False, 'message': 'Candidate not found'}), 404

    # Allow recruiter who owns candidate OR admin from same company
    if not (
        candidate.recruiter_id == current_user.id or
        (current_user.role == 'admin' and current_user.company_id == User.query.get(candidate.recruiter_id).company_id)
    ):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    keywords = (candidate.key_skills or "").split(",")
    location = candidate.preferred_location or "remote"
    remote = True

    if not any(keywords):
        keywords = ["software", "developer"]

    from universal_job_scraper import main as scraper_main

    def run_scraper_only():
        with app.app_context():
            jobs = scraper_main(
                keywords=keywords,
                location=location,
                remote=remote,
                user_id=candidate.id,
                # candidate_id=candidate.id,  # Changed back to user_id
                db=db,
                Job=Job
            )
            if jobs:
                print("✅ Scraping complete.")
                try:
                    print("⚙️ Running check.py automation...")
                    subprocess.run([sys.executable, "check.py"], check=True)
                except subprocess.CalledProcessError as e:
                    print(f"❌ Error running check.py: {e}")
            else:
                print("⚠️ No jobs scraped.")

    threading.Thread(target=run_scraper_only).start()

    return jsonify({'success': True, 'message': 'JobJarvis scraping started'})

@app.route('/pause-job-jarvis/<int:candidate_id>', methods=['POST'])
def pause_job_jarvis(candidate_id):
    # Optional: Set pause flag in DB if implementing real pause logic
    return jsonify({'success': True, 'message': 'JobJarvis paused'})

@app.route('/profile/<int:candidate_id>')
def profile_section(candidate_id):
    if 'user_id' not in session:
        return redirect('/login')

    candidate = Candidate.query.get(candidate_id)
    current_user = User.query.get(session['user_id'])

    if not candidate:
        return "Candidate not found", 404

    # ✅ Jobseeker can only see their own profile
    if current_user.role == 'jobseeker':
        my_candidate = Candidate.query.filter_by(email=current_user.email).first()
        if my_candidate and my_candidate.id == candidate_id:
            return render_template(
                'profile_section.html',
                candidate_id=candidate.id,
                full_name=candidate.name,
                email=candidate.email,
                phone=candidate.phone,
                title=candidate.title,
                resume_path=candidate.resume_path,
                preferred_location=candidate.preferred_location,
                work_authorization=candidate.work_authorization,
                key_skills=candidate.key_skills,
                employment=candidate.employment,
                education=candidate.education,
                certifications=candidate.certifications,
                projects=candidate.projects
            )
        else:
            return "Unauthorized access", 403

    # ✅ Recruiter who owns candidate OR admin from same company
    if (
        candidate.recruiter_id == current_user.id or
        (current_user.role == 'admin' and current_user.company_id == User.query.get(candidate.recruiter_id).company_id)
    ):
        return render_template(
            'profile_section.html',
            candidate_id=candidate.id,
            full_name=candidate.name,
            email=candidate.email,
            phone=candidate.phone,
            title=candidate.title,
            resume_path=candidate.resume_path,
            preferred_location=candidate.preferred_location,
            work_authorization=candidate.work_authorization,
            key_skills=candidate.key_skills,
            employment=candidate.employment,
            education=candidate.education,
            certifications=candidate.certifications,
            projects=candidate.projects
        )

    return "Unauthorized access", 403

@app.route('/profile', methods=['POST'])
def update_profile():
    candidate_id = request.form.get('candidate_id') or session.get('candidate_id')
    if not candidate_id:
        return jsonify({'error': 'Candidate not logged in'}), 401

    candidate = Candidate.query.get(int(candidate_id))
    if not candidate:
        return jsonify({'error': 'Candidate not found'}), 404

    current_user = User.query.get(session['user_id'])
    if not current_user:
        return jsonify({'error': 'User not logged in'}), 401

    # ✅ Jobseeker: can only update their own candidate record
    if current_user.role == 'jobseeker':
        my_candidate = Candidate.query.filter_by(email=current_user.email).first()
        if not my_candidate or my_candidate.id != int(candidate_id):
            return jsonify({'error': 'Unauthorized'}), 403

    # ✅ Recruiter/Admin: can update if they own candidate or are allowed as admin
    elif not (
        candidate.recruiter_id == current_user.id or
        (current_user.role == 'admin' and current_user.company_id == User.query.get(candidate.recruiter_id).company_id)
    ):
        return jsonify({'error': 'Unauthorized'}), 403

    action = request.form.get('action')

    # ✅ Handle resume removal
    if action == 'remove_resume':
        candidate.resume_path = None
        db.session.commit()
        return jsonify({'message': 'Resume removed successfully'})

    # ✅ Handle autofill from resume
    elif action == 'autofill_resume':
        if 'resume' in request.files:
            file = request.files['resume']
            if file.filename != '':
                filename = secure_filename(file.filename)
                save_path = os.path.join("resumes", filename)
                file.save(save_path)
                candidate.resume_path = save_path
                db.session.commit()

        if not candidate.resume_path:
            return jsonify({'error': 'No resume uploaded'}), 400

        parsed = parse_resume(candidate.resume_path)

        # Normalize to the career-info fields we care about
        if isinstance(parsed, dict):
            fields = {
                'key_skills': (parsed.get('key_skills') or '')[:2000],
                'employment': parsed.get('employment') or '',
                'education': parsed.get('education') or '',
                'certifications': parsed.get('certifications') or '',
                'projects': parsed.get('projects') or ''
            }
        else:
            # Fallback if parser returns raw text (string)
            text = parsed or ''
            fields = {
                'key_skills': text[:2000],
                'employment': '',
                'education': '',
                'certifications': '',
                'projects': ''
            }

        # Save ONLY career info into DB first
        candidate.key_skills     = fields['key_skills']
        candidate.employment     = fields['employment']
        candidate.education      = fields['education']
        candidate.certifications = fields['certifications']
        candidate.projects       = fields['projects']
        db.session.commit()

        # Then return the saved values for the frontend to display
        return jsonify({
            'message': 'Resume parsed and saved',
            'profile_data': fields
        })

    # ✅ Handle resume upload only (no action passed)
    if 'resume' in request.files:
        file = request.files['resume']
        if file.filename != '':
            filename = secure_filename(file.filename)
            save_path = os.path.join("resumes", filename)
            file.save(save_path)
            candidate.resume_path = save_path
            db.session.commit()
            return jsonify({'message': 'Resume uploaded successfully'})

    # ✅ Handle profile section updates
    fields = [
        'first_name', 'last_name', 'email', 'phone', 'title',
        'work_authorization', 'preferred_location', 'key_skills',
        'employment', 'education', 'certifications', 'projects'
    ]

    for field in fields:
        if field in request.form:
            value = request.form.get(field)
            if field in ['first_name', 'last_name']:
                name_parts = candidate.name.split()
                if field == 'first_name':
                    candidate.name = f"{value} {name_parts[1]}" if len(name_parts) > 1 else value
                else:
                    candidate.name = f"{name_parts[0]} {value}" if name_parts else value
            else:
                setattr(candidate, field, value)

    db.session.commit()
    return jsonify({'message': 'Profile updated successfully'})

@app.route('/autofill-resume/<int:candidate_id>', methods=['POST'])
def autofill_resume(candidate_id):
    candidate = Candidate.query.get(candidate_id)
    if not candidate or candidate.recruiter_id != session['user_id']:
        return jsonify({'message': 'Unauthorized'}), 403
    if not candidate.resume_path:
        return jsonify({'message': 'No resume uploaded'}), 400

    parsed_text = parse_resume(candidate.resume_path)
    candidate.key_skills = parsed_text[:500]  # You can later improve this logic
    db.session.commit()
    return jsonify({'message': 'Resume data extracted and saved'})

@app.route('/get-applications')
def get_applications():
    candidate_id = request.args.get('candidate_id') or session.get('candidate_id')
    if not candidate_id:
        return jsonify({'error': 'Candidate not logged in'}), 401

    candidate_id = int(candidate_id)
    jobs = Job.query.filter_by(candidate_id=candidate_id).order_by(Job.created_at.desc()).all()


    applications = []
    for job in jobs:
        applications.append({
            'title': job.title,
            'location': job.location,
            'company': job.company,
            'created_at': job.created_at.strftime('%Y-%m-%d %H:%M') if job.created_at else '',
            'status': job.status or 'Applied',  # Default if missing
            'comment': job.comment or '',
            'interview_status': job.interview_status or '',
            'final_verdict': job.final_verdict or ''
        })

    return jsonify({'status': 'success', 'applications': applications})

@app.route('/get-jobs')
def get_jobs():
    candidate_id = request.args.get('candidate_id') or session.get('candidate_id')
    if not candidate_id:
        return jsonify({'status': 'error', 'message': 'Candidate not logged in'}), 401

    candidate_id = int(candidate_id)
    jobs = Job.query.filter_by(candidate_id=candidate_id).order_by(Job.created_at.desc()).all()

    jobs_data = []
    for job in jobs:
        jobs_data.append({
            'title': job.title,
            'company': job.company,
            'location': job.location,
            'link': job.link,
            'source': job.source,
            'created_at': job.created_at.strftime('%Y-%m-%d %H:%M') if job.created_at else '',
            'status': job.status,
            'comment': job.comment,
            'interview_status': job.interview_status,
            'final_verdict': job.final_verdict
        })

    return jsonify({'status': 'success', 'jobs': jobs_data})


@app.route("/process-jobs", methods=["POST"])
def process_jobs():
    candidate_id = request.form.get("candidate_id") or session.get("candidate_id")
    if not candidate_id:
        return jsonify({"error": "Candidate not logged in"}), 401

    candidate_id = int(candidate_id)

    # Get all jobs for the candidate
    jobs = Job.query.filter_by(candidate_id=candidate_id).order_by(Job.created_at.desc()).all()
    if not jobs:
        return jsonify({"error": "No jobs found for candidate"}), 404

    for job in jobs:
        if not job.link:
            continue

        link_lower = job.link.lower()
        if "indeed" in link_lower:
            threading.Thread(target=run_indeed_for_candidate, args=(candidate_id, app)).start()
        elif "linkedin" in link_lower:
            threading.Thread(target=run_linkedin_for_candidate, args=(candidate_id, app)).start()
        else:
            app.logger.warning(f"Unknown source for candidate {candidate_id}: {job.link}")

    return jsonify({"status": "processing_started"})


if __name__ == '__main__':
    app.run(debug=True, port=5002)